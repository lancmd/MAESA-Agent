#!/usr/bin/env python3
"""Create the Wanbei six-city mining-area ecosystem-service report.

The document follows the chapter logic used by the two supplied theses while
keeping all numerical claims tied to the current project outputs.  It creates
three-line tables, places the generated maps beside the relevant discussion,
and records validation limitations instead of promoting unverified evidence to a
completed accuracy assessment.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


YEARS = [2005, 2010, 2015, 2020, 2025]
CLASS_ORDER = ["沉陷积水", "自然水体", "建设用地", "耕地", "林地", "草地"]
SCENARIOS = ["ND", "UD", "EP", "RE"]
SCENARIO_CN = {"ND": "自然发展", "UD": "城镇发展", "EP": "生态保护", "RE": "资源开采"}
CN_NUMERALS = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "七", 8: "八", 9: "九", 10: "十"}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def fmt(value: float, digits: int = 2) -> str:
    return f"{value:,.{digits}f}"


def pct(a: float, b: float) -> float:
    return (a / b - 1.0) * 100.0 if b else math.nan


def numeric_field(row: dict[str, str], *names: str) -> float:
    """Return a numeric value from either English or Chinese summary columns."""
    for name in names:
        value = row.get(name)
        if value not in (None, ""):
            return float(value)
    raise KeyError(f"none of {names} found in row with fields {sorted(row)}")


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:" + key), str(value))


def three_line_table(document: Document, title: str, headers: list[str], rows: list[list[str]], note: str | None = None):
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_before = Pt(5)
    caption.paragraph_format.space_after = Pt(2.5)
    run = caption.add_run(title)
    run.bold = True
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(text)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for record in rows:
        cells = table.add_row().cells
        for index, text in enumerate(record):
            cells[index].text = str(text)
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    none = {"val": "nil"}
    top = {"val": "single", "sz": "12", "color": "000000"}
    mid = {"val": "single", "sz": "6", "color": "000000"}
    bottom = {"val": "single", "sz": "12", "color": "000000"}
    for r_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell, top=top if r_index == 0 else none,
                            bottom=mid if r_index == 0 else bottom if r_index == len(table.rows) - 1 else none,
                            left=none, right=none, insideH=none, insideV=none)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(9.5)
                    if r_index == 0:
                        run.bold = True
    if note:
        p = document.add_paragraph("注：" + note)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(8.5)
    return table


def set_run_font(run, size=12, bold=False, east="宋体", west="Times New Roman") -> None:
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.bold = bold


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.strip())
    set_run_font(run, 12)


# Word's native OMML math objects are the portable equivalent of the
# equation objects produced by MathType.  Building them explicitly keeps
# subscripts, superscripts and fractions editable in Word instead of leaving
# a plain-text formula that changes appearance across machines.
def _m(tag: str):
    return OxmlElement(f"m:{tag}")


def _m_run(text: str):
    run = _m("r")
    t = _m("t")
    t.text = text
    run.append(t)
    return run


def _m_sub(base, sub):
    node = _m("sSub")
    e = _m("e"); e.append(base)
    s = _m("sub"); s.append(sub)
    node.extend([e, s])
    return node


def _m_sup(base, sup):
    node = _m("sSup")
    e = _m("e"); e.append(base)
    s = _m("sup"); s.append(sup)
    node.extend([e, s])
    return node


def _m_frac(num, den):
    node = _m("f")
    n = _m("num"); n.append(num)
    d = _m("den"); d.append(den)
    node.extend([n, d])
    return node


def _m_group(parts: list):
    """Group several OMML runs as one numerator/denominator expression."""
    node = _m("box")
    e = _m("e")
    for part in parts:
        e.append(part)
    node.append(e)
    return node


def add_formula(document: Document, number: str, parts: list) -> None:
    """Insert a centered, editable OMML equation with a thesis-style label."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    math_para = _m("oMathPara")
    math = _m("oMath")
    for part in parts:
        math.append(part)
    math_para.append(math)
    p._p.append(math_para)
    label = p.add_run(f"   （{number}）")
    set_run_font(label, 10.5, False, "宋体", "Times New Roman")


def add_figure(document: Document, path: Path, caption: str, width_cm: float = 15.2) -> None:
    if not path.is_file():
        add_body(document, f"图件缺失记录：预期图件“{caption}”未在成果目录中找到，正文不以空白占位图替代。")
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    c = document.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.keep_with_next = True
    c.paragraph_format.space_after = Pt(6)
    run = c.add_run(caption)
    set_run_font(run, 10.5)


def _font(size: int, bold: bool = False):
    """Use a Windows CJK font for generated diagrams, with safe fallbacks."""
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for path in candidates:
        if path.is_file():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def build_workflow_figure(path: Path) -> None:
    """Create a clean bilingual workflow diagram for the methods chapter."""
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1900, 980
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    title_font = _font(48, True)
    head_font = _font(24, True)
    body_font = _font(22)
    small_font = _font(20)
    d.text((w // 2, 42), "皖北六市矿区生态系统服务评估流程 / MAESA workflow",
           fill="#17324d", font=title_font, anchor="ma")
    nodes = [
        ("数据准备\nData preparation", "影像、逐期ROI、六市/矿区边界\nDEM、气候、土壤、威胁与沉陷证据", "#e7f0f7"),
        ("分类与统一\nClassification & harmonization", "ENVI监督分类 → 六类编码\n30 m主网格、共同支撑区、单位检查", "#eaf4e2"),
        ("土地利用模拟\nPLUS–CARS scenarios", "2020—2025扩张驱动\nND / UD / EP / RE四情景（RE含沉陷/工作面）", "#fff2d8"),
        ("生态服务计算\nInVEST services", "Carbon：Mg C\nWater Yield：mm / m³\nHabitat Quality：0—1", "#fbe7e7"),
        ("综合评价与输出\nAssessment & reporting", "全时期min–max + AHP\n图件、三线表、转移矩阵、桑基图\n现场核查与复现清单", "#ece8f7"),
    ]
    x0, bw, bh, gap = 85, 315, 380, 55
    y = 230
    for i, (head, body, fill) in enumerate(nodes):
        x = x0 + i * (bw + gap)
        d.rounded_rectangle((x, y, x + bw, y + bh), radius=24, fill=fill, outline="#2e4a62", width=4)
        d.text((x + bw / 2, y + 64), head, fill="#17324d", font=head_font, anchor="mm", align="center")
        d.line((x + 34, y + 112, x + bw - 34, y + 112), fill="#9aa9b6", width=2)
        d.multiline_text((x + bw / 2, y + 205), body, fill="#23313f", font=body_font,
                         anchor="mm", align="center", spacing=15)
        d.text((x + bw / 2, y + bh - 34), f"Step {i + 1}", fill="#557086", font=small_font, anchor="mm")
        if i < len(nodes) - 1:
            ax = x + bw + 12
            # Keep connectors in the header band so they never cross the
            # explanatory text inside a node.
            ay = y + 72
            end = x + bw + gap - 16
            d.line((ax, ay, end, ay), fill="#405b70", width=7)
            d.polygon([(end, ay), (end - 22, ay - 15), (end - 22, ay + 15)], fill="#405b70")
    d.rounded_rectangle((85, 735, w - 85, 890), radius=20, fill="#f6f8fa", outline="#9aa9b6", width=3)
    d.multiline_text((w / 2, 812), "证据链：原始数据 → 可追溯中间栅格 → 模型输出 → 统计与制图 → 现场核查\n"
                     "Evidence chain: raw inputs → harmonized grids → model outputs → tables & maps → field inspection",
                     fill="#2f4858", font=body_font, anchor="mm", align="center", spacing=10)
    im.save(path, quality=95)


def build_field_photo_panel(path: Path, photo_dir: Path) -> None:
    """Compose the three supplied field photos into a report-ready panel."""
    files = [
        (photo_dir / "20260902-172602.jpg", "南湖：湿地恢复型"),
        (photo_dir / "20260902-172627.jpg", "临涣：农业复垦与沉陷治理型"),
        (photo_dir / "e29e6448-ebfb-45ed-84e4-ae583fa3dac4-File", "潘集：光伏或新能源利用型"),
    ]
    existing = [(p, label) for p, label in files if p.is_file()]
    if not existing:
        return
    panel_w, panel_h, label_h = 620, 480, 90
    canvas = Image.new("RGB", (panel_w * len(existing), panel_h + label_h), "white")
    d = ImageDraw.Draw(canvas)
    label_font = _font(27, True)
    for i, (photo, label) in enumerate(existing):
        img = Image.open(photo).convert("RGB")
        img.thumbnail((panel_w - 20, panel_h - 20), Image.Resampling.LANCZOS)
        x = i * panel_w + (panel_w - img.width) // 2
        y = (panel_h - img.height) // 2
        canvas.paste(img, (x, y))
        d.rectangle((i * panel_w, panel_h, (i + 1) * panel_w, panel_h + label_h), fill="#eef3f6")
        d.text((i * panel_w + panel_w / 2, panel_h + label_h / 2), label, fill="#1f3547", font=label_font, anchor="mm")
        if i:
            d.line((i * panel_w, 0, i * panel_w, panel_h + label_h), fill="#ffffff", width=5)
    path.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(path, quality=94)


def add_extended_analysis(document: Document, h: dict, data: dict, s: dict) -> None:
    """Add the long-form discussion requested for a thesis-style report.

    The main chapters present the reproducible results.  This chapter adds the
    interpretive layer normally found in a thesis: regional setting, process
    mechanisms, cross-service comparison, scenario meaning and application.
    Every numeric statement is derived from the current common-support series.
    """
    service_order = "、".join(SCENARIO_CN[code] for code in sorted(SCENARIOS, key=lambda code: s[code]["service_mean"], reverse=True))
    habitat_order = "、".join(SCENARIO_CN[code] for code in sorted(SCENARIOS, key=lambda code: s[code]["habitat_mean"], reverse=True))
    sections = [
        ("11.1 五期土地利用变化的区域结构", [
            f"共同支撑区内，耕地面积由2005年的{data['lulc_area'][2005]['耕地']:.2f} hm²变化到2025年的{data['lulc_area'][2025]['耕地']:.2f} hm²，始终是占比最大的地类。建设用地在2005年为{data['lulc_area'][2005]['建设用地']:.2f} hm²，2020年降至{data['lulc_area'][2020]['建设用地']:.2f} hm²，2025年为{data['lulc_area'][2025]['建设用地']:.2f} hm²；这一序列呈现出先降后稳的阶段性，而不是单调扩张。沉陷积水从{data['lulc_area'][2005]['沉陷积水']:.2f} hm²增加到2020年的{data['lulc_area'][2020]['沉陷积水']:.2f} hm²，2025年回落到{data['lulc_area'][2025]['沉陷积水']:.2f} hm²，说明水陆边界和治理利用方式可能在不同阶段发生调整。",
            "从空间格局看，矿区并非连续的单一采矿地表，而是被村庄、农田、道路和河沟切割的多斑块系统。耕地是矿区周边最重要的背景地类，沉陷积水往往在其内部或边缘形成；建设用地则沿交通和居民点分布，成为威胁暴露与服务短板的潜在叠加区。林地和草地面积虽然相对较小，却承担岸带缓冲、斑块连接和高碳密度补偿等功能，不能因为面积小而在治理排序中被忽略。",
            "五期面积差异应按“结构信号”和“分类信号”两层阅读。结构信号包括耕地长期主体、沉陷积水在采煤集中区出现、建设用地与交通走廊相邻等稳定关系；分类信号则表现为自然水体、林地和草地在相邻期的大幅跳变。报告将六类编码固定为沉陷积水、自然水体、建设用地、耕地、林地和草地，后续若重新分类，只需替换栅格和统计表，不改变碳、水、生境模型的接口。",
        ]),
        ("11.2 沉陷积水与矿区治理逻辑", [
            f"沉陷积水是本研究区别于一般土地利用评价的关键对象。2005—2025年其面积序列为{data['lulc_area'][2005]['沉陷积水']:.2f}、{data['lulc_area'][2010]['沉陷积水']:.2f}、{data['lulc_area'][2015]['沉陷积水']:.2f}、{data['lulc_area'][2020]['沉陷积水']:.2f}和{data['lulc_area'][2025]['沉陷积水']:.2f} hm²。面积在2015—2020年快速增加后有所回落，可能对应沉陷扩展、季节水位、复垦填筑或水体边界识别差异。只有将沉陷云图、工作面、历史影像和水系叠加，才能区分这些机制。",
            "治理上应把积水区分成至少三类：第一类是水质较好、岸带稳定且与河网或湿地具有连通潜力的生态恢复单元；第二类是与耕地、村庄和排灌设施高度交叠的生产安全单元；第三类是受矿业设施、道路或污染源影响的风险控制单元。三类单元的目标分别是提高湿地与生境功能、保障土地利用和排灌安全、控制污染与地质灾害风险。统一按“积水面积越大越好”处理会掩盖这种差异。",
            "南湖、临涣和潘集现场照片分别对应湿地恢复、农业复垦与沉陷治理、光伏或新能源利用三种路径。它们说明同样的沉陷水面可以被赋予不同的管理目标：南湖强调岸带、游憩与湿地生态；临涣强调耕地生产、排灌和沉陷稳定；潘集强调水面利用与设施布局的兼容。报告中的模型结果用于识别空间优先序，照片用于补充场景解释，后续应以水深、水质、底泥、植被和工程安全调查完成定量闭环。",
        ]),
        ("11.3 碳储量变化与地类贡献", [
            f"统一碳密度表和共同支撑区下，总碳储量由2005年的{h[2005]['carbon_mg_c']/1e6:.3f}×10^6 Mg C变化到2010年的{h[2010]['carbon_mg_c']/1e6:.3f}×10^6 Mg C，2015年升至{h[2015]['carbon_mg_c']/1e6:.3f}×10^6 Mg C，2020年降至{h[2020]['carbon_mg_c']/1e6:.3f}×10^6 Mg C，2025年为{h[2025]['carbon_mg_c']/1e6:.3f}×10^6 Mg C。全期振幅约为{(max(v['carbon_mg_c'] for v in h.values())-min(v['carbon_mg_c'] for v in h.values()))/1e6:.3f}×10^6 Mg C，明显小于水源供给总量的相对波动。",
            "碳储量变化的首要控制量是地类面积与单位面积碳密度的乘积。耕地面积大，因而即使单位碳密度不是最高，也对总量具有稳定贡献；林地单位碳密度较高，在面积变化较小的情况下也可能显著影响局地碳储量；沉陷积水和自然水体的碳库则应结合水体、底泥和水生植被分层估算，不能简单套用陆地植被密度。",
            "沉陷积水复合碳库为区域碳储量提供了除陆地地类外的补充视角。水体库容决定溶解碳和颗粒碳的容纳空间，水生植被提供短周期生物量，底泥则可能形成相对稳定的有机碳库；但沉积、再悬浮、藻类生产和水质变化会改变碳的存留时间。工程上宜先建立典型水体的水深—面积—底泥厚度样线，再将参数回填到复合碳库模型。",
        ]),
        ("11.4 水源供给与气候背景", [
            f"五期平均产水深分别为{h[2005]['water_mean_mm']:.2f}、{h[2010]['water_mean_mm']:.2f}、{h[2015]['water_mean_mm']:.2f}、{h[2020]['water_mean_mm']:.2f}和{h[2025]['water_mean_mm']:.2f} mm，总量分别为{h[2005]['water_m3']/1e6:.3f}、{h[2010]['water_m3']/1e6:.3f}、{h[2015]['water_m3']/1e6:.3f}、{h[2020]['water_m3']/1e6:.3f}和{h[2025]['water_m3']/1e6:.3f}×10^6 m³。2010年为低值期，2005年和2020年为相对高值期，2025年回落。该变化首先反映降水和参考蒸散背景，其次才是土地利用和土壤持水条件的调节。",
            "Budyko框架的优点是把降水、潜在蒸散和下垫面参数放入同一水量平衡。对于矿区而言，沉陷造成的微地形汇水、人工排水、地下水位变化和水体蒸发会使实际过程偏离自然下垫面假设。因此，年产水量图适合回答“哪些网格的产水潜力较高”，不适合直接回答“有多少水可以安全取用”。水质、水位、工程调度和污染负荷必须单独纳入决策。",
            "在治理应用中，高产水区并不必然是优先保护区：若高产水与污染源、建设用地或陡岸重叠，可能需要先进行风险控制；低产水区也不必然是低价值区，稳定林草和农田土壤可能具有重要蓄水和抗旱功能。建议将年产水深与河网、沉陷水体、排水工程和水质监测点叠加，形成“水量潜力—调蓄能力—环境风险”三维分区。",
        ]),
        ("11.5 生境质量与威胁暴露", [
            f"生境质量均值从2005年的{h[2005]['habitat_mean']:.4f}上升至2010年的{h[2010]['habitat_mean']:.4f}、2015年的{h[2015]['habitat_mean']:.4f}和2020年的{h[2020]['habitat_mean']:.4f}，2025年为{h[2025]['habitat_mean']:.4f}。这一序列表明，在当前地类适宜性、道路/铁路/建设/矿业威胁权重和距离衰减参数下，2020年形成相对高值，2025年出现调整。它是模型代理指标，不能等同于物种数量或群落完整性。",
            "空间上，生境质量更容易在稳定林草斑块、水体岸带和远离高强度建设的地段形成高值，而在道路交叉、矿业设施、居民点外缘和斑块边缘形成低值或退化带。威胁的最大作用距离决定了低值带宽度，敏感性表决定了相同威胁在耕地、水体、林地和草地上的差异。参数改变可能导致局地排序变化，因此应在治理前开展权重、距离和Z参数的敏感性分析。",
            "矿区生境治理不应只增加绿地面积。对于高质量斑块，应限制道路和建设切割，维护其内部连续性；对于中等质量的水体—林草过渡带，应优先建设缓冲带、浅滩和生态廊道；对于多威胁叠加的低值带，应先控制扰动源和污染，再实施植被恢复。将生境质量图与沉陷云图和工作面叠加，可以识别采动风险可能切断生态连通性的关键位置。",
        ]),
        ("11.6 综合生态服务的协同与权衡", [
            f"在共同支撑区和统一min-max范围下，综合指数由2005年的{h[2005]['service_mean']:.4f}降至2010年的{h[2010]['service_mean']:.4f}，2015年回升至{h[2015]['service_mean']:.4f}，2020年达到{h[2020]['service_mean']:.4f}，2025年为{h[2025]['service_mean']:.4f}。阶段变化与碳储量、水源供给、生境质量并不同步，说明多服务之间存在时间错位和空间权衡。",
            "AHP权重把碳储量放在首位，因而碳库变化对综合指数的影响更大；水源供给的波动幅度较大，可能在丰水年提升综合指数；生境质量虽然权重较低，却是识别矿业威胁和生态连通性的关键。线性加权便于比较，但存在补偿效应：一个服务的下降可能被另一个服务的上升抵消。因此报告中始终同时给出三项单项服务，不用综合指数替代任何单项诊断。",
            "建议把综合服务分区分成“协同保护、单项短板、冲突权衡、综合修复”四类。协同保护区在三项服务上均保持较高水平，应减少新增扰动；单项短板区需要先确定短板指标及其成因；冲突权衡区应采用分期、分层工程，避免只提高某项指标；综合修复区则需要把沉陷治理、林草恢复、耕地保护和威胁控制组合实施。",
        ]),
        ("11.7 2026年四情景的机制解释", [
            f"四情景服务统计显示，ND、UD、EP、RE的碳储量分别为{s['ND']['carbon_mg_c']/1e6:.3f}、{s['UD']['carbon_mg_c']/1e6:.3f}、{s['EP']['carbon_mg_c']/1e6:.3f}和{s['RE']['carbon_mg_c']/1e6:.3f}×10^6 Mg C；平均产水深分别为{s['ND']['water_mean_mm']:.2f}、{s['UD']['water_mean_mm']:.2f}、{s['EP']['water_mean_mm']:.2f}和{s['RE']['water_mean_mm']:.2f} mm。四情景沿用2025年气候背景，因此差异主要来自土地利用空间配置和参数框架，而不是新增的气候变化信号。",
            f"当前配置下，综合生态服务指数由高到低为{service_order}，生境质量由高到低为{habitat_order}。排序差异说明情景名称不等于服务效果，土地需求、转换矩阵、限制区和威胁栅格共同决定最终结果。规划解释时应先核查规则是否落实，再比较服务差异。",
            "RE情景叠加沉陷深度和工作面证据，适合用于识别资源开采条件下的风险敏感区，而不是作为确定性开采预测。对RE中耕地减少、沉陷积水增加或生境质量变化明显的网格，应回看工作面推进、沉陷范围、水系连通和土地需求约束；对四情景均稳定的网格，则可作为长期保护或低冲突利用的候选区。",
        ]),
        ("11.8 典型矿区案例与现场照片的互证", [
            "南湖湿地恢复型案例体现了沉陷水体由地质环境问题转变为生态空间的可能路径。评价重点应包括水体面积和水深变化、岸带植被连续性、与周边林草斑块的连通性、游憩设施对核心生境的影响以及水质风险。对于这类矿区，综合生态服务提升不应只依赖水面扩大，而应关注岸带缓冲和水生生物栖息条件是否同步改善。",
            "临涣农业复垦与沉陷治理型案例表明，农田恢复与沉陷风险控制需要同步推进。复垦地的覆土厚度、土壤有机质、排灌条件和作物长势决定耕地能否稳定承载生产功能；若沉陷仍在发展，则短期面积恢复可能再次转为积水或低效利用。报告建议将耕地面积、沉陷深度、工作面距离、产水深和碳储量纳入同一矿区台账，按年度更新。",
            "潘集光伏或新能源利用型案例说明水面利用会引入新的空间权衡。漂浮光伏可提高沉陷水体的利用效率，但设施覆盖率、锚固安全、岸带通行、水体光照和鸟类活动均可能影响生境质量。对于该类矿区，应在情景模拟中增加水面设施占用和生态缓冲约束，并通过现场测量校正水体碳库、蒸发和威胁距离参数。",
        ]),
        ("11.9 结果复现与论文写作衔接", [
            "论文式结果需要同时呈现方法、数据、公式、图件和表格，而不是只展示最终数值。本报告保留了六类LULC编码、共同支撑区、栅格分辨率、碳密度表、气候与土壤输入、威胁敏感性表、AHP权重和情景需求清单，使读者可以从原始数据一路追溯到模型结果。图件标题采用“年份/情景＋指标＋空间分布”的形式，表格统一使用hm²、Mg C、mm、m³和0—1等单位。",
            "对于实际论文，建议在正文方法章节中先说明共同边界和分类体系，再分节讨论土地利用、碳储量、水源供给、生境质量和综合服务；在结果章节中每个小节遵循“变化描述—空间特征—机制解释—治理含义”的顺序；在讨论章节中把沉陷、复垦、气候和参数不确定性放到同一框架下。这样既能保持报告的工程可读性，也符合论文对逻辑链和可重复性的要求。",
            "本报告已形成可直接用于论文结果章节的流程复现、相对比较和治理方案筛选成果。后续可围绕典型水体水文、水质、样地碳密度、物种样点和PLUS历史回代开展参数本地化；新增资料后在不改变主网格和编码的前提下增量重算，并同步更新图、表和运行清单。",
        ]),
        ("11.10 成果应用与后续更新", [
            "成果可直接服务三类工作：一是矿区生态修复规划，通过单项服务和综合服务图识别优先修复单元；二是采煤沉陷动态监管，通过工作面、沉陷云图、沉陷积水分类和RE情景识别风险带；三是论文与项目验收，通过三线表、年度图件、转移矩阵、桑基图和流程图形成完整证据链。三类应用共用同一套分类编码和数据清单，避免重复制图和口径分裂。",
            "建议建立年度更新节奏：每年更新遥感影像、工作面和沉陷云图；在关键开采节点更新土地利用与四项生态服务；在修复工程完成后补充现场照片和样点；在模型参数或软件版本变化时保存版本、哈希和变更说明。对于水体、耕地、林草和建设用地的突变，应先做空间复核，再决定是否进入政策评估或工程绩效统计。",
            "当本地实测数据逐步补齐后，可进一步开展不确定性传播、权重敏感性、典型矿区对照和成本—效益分析，把当前的相对服务指数升级为带置信区间的决策支持结果。对外发布时保留匿名示例项目、软件版本和运行日志，对内则保留原始影像、ROI、参数表和现场调查记录，从而兼顾数据安全、科研诚信和长期维护。",
        ]),
    ]
    sections.extend([
        ("11.11 图件与统计表的联合判读", [
            "论文中的空间图件并不是统计表的装饰，而是把数量变化落实到地理位置的证据。面积表可以说明某一地类增加或减少，转移矩阵可以说明编码变化的方向，地图则进一步回答变化发生在何处、是否集中在工作面周边、是否沿道路或水系扩展。将三者联合使用，能够把“面积变化”转化为“矿区过程线索”：例如沉陷积水增加且集中在工作面下方时，应检查沉陷深度和水位；建设用地增加且沿交通廊道分布时，应检查扩张驱动和缓冲范围。",
            "碳储量图、水源供给图、生境质量图和综合服务图需要按同一空间位置叠加阅读。某一网格碳储量较高但生境质量偏低，可能说明存在高碳作物地与交通威胁叠加；某一网格年产水深较高但综合服务并不高，可能与威胁暴露或低碳密度地类有关。因而报告不采用“单幅图下结论”的方式，而是通过图组、时间图和统计表共同构建解释链。",
            "图件标题、图例和单位在报告中保持一致，是保证结果可读性的基础。六类土地利用图统一采用沉陷积水、自然水体、建设用地、耕地、林地和草地；碳储量采用Mg C或10^6 Mg C，水源供给采用mm和10^6 m³，生境质量与综合服务采用0—1。这样的单位契约让读者能够直接对照地图、时间曲线和三线表，减少不同章节之间的误读。",
        ]),
        ("11.12 阶段性变化的过程解释", [
            "2005—2010年是研究期的初始调整阶段，沉陷积水、建设用地和耕地之间的转换构成主要空间背景；2010—2015年耕地面积增加、沉陷积水继续扩展，可能对应复垦、分类规则调整和开采阶段变化的共同作用；2015—2020年沉陷积水达到较高水平，林草面积和生境质量也出现阶段性变化；2020—2025年部分水体和林草面积回落，综合服务指数随之调整。这里的阶段划分用于组织结果，不把每个拐点直接归因于单一工程或政策。",
            "阶段性变化还反映了矿区系统的滞后效应。采煤扰动发生后，沉降、积水、复垦和植被恢复并不在同一年完成；碳库恢复通常存在多年时滞，水源供给受气候年景影响立即波动，生境质量则对威胁距离和斑块连通性较敏感。把所有指标放在同一年度比较，只能得到状态差异；要识别过程，还需要补充工作面推进时间、治理工程时间和连续遥感序列。",
            "因此，报告建议在后续更新中增加“开采节点—地类转换—服务响应”三联表。对于每个典型矿区，记录工作面开始和结束时间、沉陷范围变化、积水面积、水生植被和底泥情况，以及碳、水、生境和综合指数的同步变化。这样可以把区域尺度的趋势分析下沉到矿区过程，形成更接近论文讨论章节的证据。",
        ]),
        ("11.13 参数来源与本地化路线", [
            "当前碳密度、Kc、根系深度、PAWC、Z、威胁权重和敏感性表优先采用用户提供资料、论文参考值和公开数据。这样做保证了流程能够运行，也使五期与四情景使用同一套参数，便于相对比较。但公开或文献参数反映的是区域平均条件，不一定完全适合皖北六市不同土壤、植被和矿区工程状态，实际应用时应建立本地化校准路线。",
            "碳密度本地化可从耕地、林地、草地、沉陷湿地和底泥五类样地开始，分别测定地上生物量、地下根系、土壤有机碳和底泥有机碳；水源供给校准可利用流量站、排水泵站、水位和降水资料对产水量进行季节或年度约束；生境质量校准可采用鸟类、两栖类、植物样方或红外相机的出现数据检验高低值区。",
            "参数校准不必一次性覆盖所有矿区。可以先选择南湖、临涣和潘集三类治理模式作为代表样区，建立小规模、可重复的样地和水文观测，再将校准结果通过分区参数表推广到相似矿区。这样既控制工作量，也能检验“治理模式—参数—生态服务响应”的解释框架。",
        ]),
        ("11.14 分类体系与沉陷证据的增强", [
            "六类土地利用体系的核心是把沉陷积水从自然水体中分离出来。自然水体代表河流、湖泊等相对稳定的水面，沉陷积水代表采煤沉陷或其治理形成的水体，两者在成因、岸带、污染风险和治理目标上不同。若将两类水体合并，土地利用转移矩阵会掩盖采煤扰动的方向，Carbon、Habitat Quality和RE情景也难以区分矿业影响。",
            "增强分类应采用“光谱/纹理分类＋沉陷证据约束＋空间后处理”三层策略。首先利用影像和逐期ROI形成六类初始分类；其次叠加沉陷云图、等值线、工作面及水系，约束沉陷积水的候选范围；最后利用连通域、最小斑块面积和边界平滑规则去除孤立噪声。每一步都应保存中间栅格，避免把后处理结果误认为原始分类器输出。",
            "在论文写作中，应把沉陷积水的面积变化与采煤空间证据分开表述：面积表说明分类结果，沉陷云图和工作面说明地质环境背景，二者在空间上相互支持但不能互相替代。对没有沉陷证据支撑的水体斑块，应保留为自然水体或待核查类别，不通过文字解释强行归入沉陷积水。",
        ]),
        ("11.15 三类典型矿区的对照框架", [
            "南湖、临涣和潘集构成了具有互补意义的三个典型案例：南湖强调“水体—湿地—岸带”生态恢复，临涣强调“沉陷—农田—复垦”生产治理，潘集强调“水面—新能源—设施”复合利用。三类案例可以作为区域评价结果的现场对照：如果模型把南湖识别为水体和生境协同高值区，应检查岸带实际植被和水质；如果把临涣识别为耕地和沉陷风险叠加区，应检查排灌和覆土；如果把潘集识别为水体利用和建设威胁并存区，应检查光伏设施覆盖和生态缓冲。",
            "典型矿区对照还可以帮助解释综合服务指数的差异。湿地恢复可能提升生境和水体调蓄，但不一定提升陆地碳密度；农业复垦可能维持较高耕地碳库和产水潜力，但在沉陷风险持续时生境改善有限；新能源利用可能提高土地利用效率，却引入设施遮挡、岸带切割和水鸟干扰等新的威胁。评价应以单项服务和现场事实共同判断，而不是只比较一个综合值。",
            "建议在后续论文中为每个典型矿区设置同样的图表组合：位置与底图、五期六类LULC、沉陷范围、碳储量曲线、水源供给曲线、生境质量曲线、综合服务曲线和治理建议。统一模板能让不同矿区横向可比，也能把区域尺度的复杂结果转化为工程人员易于理解的案例卡片。",
        ]),
        ("11.16 从相对评估到决策支持", [
            "当前成果已经形成从输入数据到地图、统计表和报告的完整工作流，但相对评估与决策支持之间仍有一步距离。决策需要回答治理投入放在哪里、先做什么、能达到什么目标以及如何评价成效。为此，应把综合指数和单项服务转化为可操作的优先级：高风险沉陷与低生境叠加区优先做安全与污染控制，高碳稳定斑块优先保护，耕地与沉陷边界优先做排灌和复垦，水体岸带与林草连接区优先做生态修复。",
            "四情景结果可以提供决策的“压力测试”。若某一地块在ND、UD、EP和RE四种情景下都保持低生境或高沉陷风险，说明治理具有较强的刚性需求；若只有RE情景显著变差，说明应把开采强度、工作面时序和沉陷影响范围作为约束；若不同情景在碳、水、生境之间出现权衡，则需要比较工程成本、土地权属和水质安全，而不能单看综合指数排序。",
            "最终，建议将本报告作为动态底图和年度更新模板，而不是一次性结论。随着新的遥感影像、沉陷监测、工程验收和现场样点加入，系统可以保持主网格和六类编码不变，只更新受影响年份或矿区，并重新生成表格、图件和文字。这样既满足论文复现，也支持矿区生态治理的持续管理。",
        ]),
    ])
    for title, paragraphs in sections:
        heading(document, title, 2)
        for paragraph in paragraphs:
            add_body(document, paragraph)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def display_heading(text: str, level: int) -> str:
    """Present the report hierarchy in the survey-report form used as reference.

    The source data and chapter logic remain MAESA's own.  The numbering keeps
    the Arabic chapter/subchapter contract used by the supplied investigation
    report, so a subsection always retains its parent chapter (for example
    ``3.2`` rather than a repeated “第二节”).
    """
    if level == 1:
        match = re.match(r"^(\d+)\s+(.+)$", text)
        if match:
            number, title = int(match.group(1)), match.group(2)
            return f"第{number}章 {title}"
    if level in (2, 3):
        return text
    return text


def heading(document: Document, text: str, level: int) -> None:
    p = document.add_heading(display_heading(text, level), level=level)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if re.match(r"^\d+\s+", text):
            p.paragraph_format.page_break_before = True
    elif level == 2:
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    else:
        p.paragraph_format.space_before = Pt(2.5)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in p.runs:
        if level == 1:
            set_run_font(run, 15, True, "黑体", "Times New Roman")
        elif level == 2:
            set_run_font(run, 14, True, "黑体", "Times New Roman")
        else:
            set_run_font(run, 12, True, "黑体", "Times New Roman")
        # Word's built-in heading styles may carry a blue theme colour.  The
        # reference investigation report uses plain black section headings.
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_subsection(document: Document, title: str, focus: str, evidence: str,
                   mechanism: str, implication: str, limitation: str | None = None) -> None:
    """Write a report-style analytical subsection.

    Earlier editions repeated the validation boundary after every result.  The
    revised layout keeps the main chapters focused on spatial pattern, process
    interpretation and management implication; evidence limitations are
    consolidated in Chapter 9.
    """
    heading(document, title, 2)
    paragraphs = [
        f"{focus}是认识皖北六市采煤扰动与生态响应关系的重要切入点。{evidence}。本报告将统计表、专题图、矿区边界和沉陷证据放在同一空间框架中阅读，先描述数量变化，再定位空间集聚和转化路径。",
        f"从过程机制看，{mechanism}。土地利用、地表形变、水文条件和人为扰动共同塑造生态服务格局：碳储量体现地类面积与单位碳密度的叠加，水源供给体现降水—蒸散—土壤水分平衡，生境质量体现适宜性与威胁距离衰减。",
        f"从空间组织看，不同矿区的耕地基底、沉陷积水规模、建设强度和植被条件并不相同。对这些差异进行分区解读，可以把区域均值还原为可识别的矿区单元，并解释同一时期不同矿区服务值的差异。",
        f"在治理应用上，{implication}。建议将沉陷积水、耕地连片区、城镇—矿业交叠区和稳定林草斑块作为不同的管理单元，分别组织水体安全与岸带修复、耕地保护与复垦、扰动边界管控以及生态廊道维护，使评价结果直接对应可执行的空间措施。",
    ]
    for paragraph in paragraphs:
        add_body(document, paragraph)


def first_existing(paths: list[Path]) -> Path | None:
    return next((path for path in paths if path.is_file()), None)


def load_data(root: Path, report_root: Path, maps_root: Path) -> dict:
    asset = report_root / "报告数据与图表" / "report_data.json"
    data = json.loads(asset.read_text(encoding="utf-8"))
    # Always prefer the latest common-support recomputation.  Older report
    # editions kept a report_data.json snapshot that predates the harmonized
    # 30 m mask and therefore cannot be used for the new narrative.
    hist_path = first_existing([
        root / "outputs" / "statistics_harmonized_v2_common_support" / "生态系统服务五期汇总统计_原生InVEST生境.csv",
        root / "outputs" / "statistics_harmonized_v2_common_support" / "生态系统服务五期汇总统计.csv",
    ])
    if hist_path:
        hist_rows = read_csv(hist_path)
        data["historical"] = {}
        for row in hist_rows:
            year = int(row["year"])
            data["historical"][year] = {
                "carbon_mg_c": float(row["carbon_storage_MgC"]),
                "water_m3": float(row["water_yield_volume_m3"]),
                "water_mean_mm": float(row["mean_water_yield_mm"]),
                "habitat_mean": float(row["mean_habitat_quality_index"]),
                "service_mean": float(row["mean_composite_ecosystem_service_index"]),
            }
        data["common_boundary_area_ha"] = float(hist_rows[0].get("common_boundary_area_ha", 0.0)) if hist_rows else 0.0
    else:
        data["historical"] = {int(k): v for k, v in data["historical"].items()}
        data["common_boundary_area_ha"] = float(data.get("common_boundary_area_ha", 0.0))

    area_path = first_existing([
        root / "outputs" / "statistics_harmonized_v2_common_support" / "lulc" / "共同支撑区土地利用面积统计.csv",
        root / "结果_原生InVEST重算" / "土地利用分类" / "五期六类土地利用面积统计_共同支撑区.csv",
        root / "outputs" / "classification" / "final_grid_v3" / "landuse_area_statistics.csv",
    ])
    area = read_csv(area_path) if area_path else []
    data["lulc_area"] = {year: {} for year in YEARS}
    for row in area:
        year = int(row["year"])
        if "code" in row:
            landuse = CLASS_ORDER[int(row["code"]) - 1]
        else:
            en_cn = {"subsidence_water": "沉陷积水", "natural_water": "自然水体", "built_up": "建设用地",
                     "cropland": "耕地", "forest": "林地", "grassland": "草地"}
            landuse = en_cn.get(row.get("landuse", ""), row.get("class", ""))
        if landuse in CLASS_ORDER:
            data["lulc_area"][year][landuse] = float(row["area_ha"])

    transition_path = first_existing([
        root / "outputs" / "statistics_harmonized_v2_common_support" / "transitions" / "2005_2025_土地利用转移矩阵.csv",
        root / "outputs" / "statistics" / "lulc_transition_2005_2025.csv",
    ])
    data["transition"] = read_csv(transition_path) if transition_path else []
    carbon_path = first_existing([
        root / "outputs" / "statistics_harmonized_v2_common_support" / "碳储量_分地类统计.csv",
        root / "outputs" / "statistics" / "carbon_storage_by_landuse_2005_2025.csv",
    ])
    data["carbon_by_class"] = read_csv(carbon_path) if carbon_path else []

    # Four-scenario statistics and actual demand are also refreshed from the
    # current CSVs so that the report does not silently reuse the old snapshot.
    scenario_stats = first_existing([
        report_root / "报告数据与图表" / "表_2026四情景生态系统服务统计_单位校正.csv",
        report_root / "报告数据与图表" / "表_2026四情景生态系统服务统计.csv",
    ])
    scenario_area = first_existing([
        report_root / "报告数据与图表" / "表_2026四情景实际土地利用面积.csv",
    ])
    if scenario_stats:
        data["scenario_2026"] = {}
        for row in read_csv(scenario_stats):
            vals = list(row.values())
            code = vals[0]
            if code not in SCENARIOS:
                continue
            data["scenario_2026"][code] = {
                "carbon_mg_c": float(vals[2]),
                "water_m3": float(vals[4]),
                "water_mean_mm": float(vals[6]),
                "habitat_mean": float(vals[7]),
                "service_mean": float(vals[8]),
                "lulc_area_ha": {c: 0.0 for c in CLASS_ORDER},
            }
        if scenario_area:
            for row in read_csv(scenario_area):
                vals = list(row.values())
                code, landuse, value = vals[0], vals[2], float(vals[3])
                if code in data["scenario_2026"] and landuse in CLASS_ORDER:
                    data["scenario_2026"][code]["lulc_area_ha"][landuse] = value

    typical = first_existing([
        maps_root / "分析统计图与表" / "典型矿区碳储量" / "典型矿区碳储量统计_2005_2025.csv",
        root / "结果" / "分析统计图与表" / "典型矿区碳储量" / "典型矿区碳储量统计_2005_2025.csv",
    ])
    data["typical_mines"] = read_csv(typical) if typical else []
    return data


def build(root: Path, result: Path, output: Path) -> dict:
    # Maps are retained in the result package, while the editable report and
    # its small CSV/figure assets live beside the DOCX.  Keeping these roots
    # separate makes a report rebuild safe after users organise their results.
    maps_root = result if result.is_dir() else (root / "结果_重绘")
    if not maps_root.is_dir():
        maps_root = root / "结果"
    report_root = output.parent
    data = load_data(root, report_root, maps_root)
    s = data.get("scenario_2026", {})
    service_order_cn = "、".join(SCENARIO_CN[code] for code in sorted(SCENARIOS, key=lambda code: s[code]["service_mean"], reverse=True))
    habitat_order_cn = "、".join(SCENARIO_CN[code] for code in sorted(SCENARIOS, key=lambda code: s[code]["habitat_mean"], reverse=True))
    result = maps_root
    assets = report_root / "报告数据与图表"
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.8)
    section.header_distance = Cm(1.3)
    section.footer_distance = Cm(1.3)
    add_page_number(section.footer.paragraphs[0])

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(0.85)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(78)
    r = title.add_run("安徽省皖北六市矿区")
    set_run_font(r, 20, True, "黑体")
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("生态系统服务综合评估报告")
    set_run_font(r, 21, True, "黑体")
    third = document.add_paragraph()
    third.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = third.add_run("土地利用变化、情景模拟与生态服务评价")
    set_run_font(r, 15, False, "宋体")
    document.add_paragraph("\n\n\n")
    for label, value in [("研究范围", "皖北六市矿区"), ("评价时段", "2005—2025年及2026年四情景"),
                         ("空间基准", "EPSG:32650，30 m分析网格"), ("完成时间", "2026年9月2日")]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(f"{label}：{value}")
        set_run_font(rr, 12)
    document.add_page_break()

    heading(document, "摘要", 1)
    h = data["historical"]
    summary = (
        f"本报告以皖北六市矿区为对象，按照土地利用动态—生态系统服务评估—情景模拟—综合评价的证据链，"
        f"整合2005、2010、2015、2020和2025年土地利用成果，以及2026年自然发展（ND）、城镇发展（UD）、生态保护（EP）和资源开采（RE）四种PLUS情景。"
        f"统一采用EPSG:32650、30 m分析主网格，将沉陷积水与自然水体分列，并以建设用地、耕地、林地和草地构成增强六类体系。"
        f"InVEST Carbon、Annual Water Yield和Habitat Quality分别用于碳储量、年产水量（水源供给潜力）和生境质量指数估算，三项指标经全时期统一min-max标准化后，"
        f"按AHP权重0.636986、0.258285和0.104729合成综合生态系统服务指数，判断矩阵CR为0.0332。结果显示，"
        f"现有期别土地利用图与固定碳密度组合后，碳储量结果由2005年的{h[2005]['carbon_mg_c']/1e6:.3f}×10^6 Mg C到2015年的{h[2015]['carbon_mg_c']/1e6:.3f}×10^6 Mg C，"
        f"2025年为{h[2025]['carbon_mg_c']/1e6:.3f}×10^6 Mg C；年产水量（水源供给潜力）在{min(v['water_m3'] for v in h.values())/1e6:.3f}—{max(v['water_m3'] for v in h.values())/1e6:.3f}×10^6 m³之间波动。"
        f"生境质量指数在2020年为{h[2020]['habitat_mean']:.4f}，综合生态服务指数在2010年为{h[2010]['service_mean']:.4f}、2020年为{h[2020]['service_mean']:.4f}。"
         f"2026年当前参数配置下，综合生态服务指数排序为{service_order_cn}，生境质量排序为{habitat_order_cn}。情景名称不能替代政策效果判断，需结合土地需求、转换规则和空间约束逐项审查。"
        f"据此，报告提出以沉陷积水治理、耕地保护、林草斑块维持、矿业扰动管控和生态服务协同提升为主线的分区建议。"
    )
    add_body(document, summary)
    p = document.add_paragraph()
    rr = p.add_run("关键词：皖北六市；采煤沉陷；土地利用变化；PLUS；InVEST；碳储量；水源供给；生境质量；生态系统服务")
    set_run_font(rr, 11, True)

    heading(document, "目录", 1)
    toc_items = [
        "1 绪论", "2 研究区、数据与技术路线", "3 土地利用格局、转移及2026年情景模拟",
        "4 碳储量动态与沉陷积水复合碳库", "5 年产水量（水源供给潜力）评估", "6 生境质量时空变化",
        "7 综合生态系统服务评价", "8 2026年四情景生态效应与治理分区", "9 综合结果判读与治理分区", "10 结论与建议",
        "11 研究区案例与方法扩展分析"
    ]
    for item in toc_items:
        p = document.add_paragraph(display_heading(item, 1))
        p.paragraph_format.left_indent = Cm(0.7)
        set_run_font(p.runs[0], 11)
    document.add_page_break()

    # Chapter 1
    heading(document, "1 绪论", 1)
    add_subsection(document, "1.1 研究背景与问题提出", "高潜水位煤矿区的土地—水—碳—生境耦合问题",
                   "采煤沉陷形成积水斑块，矿业和城镇建设改变地表覆盖，修复工程又促使部分扰动地向耕地、林地和草地转换",
                   "沉陷改变微地形和汇流条件，土地利用转换进一步改变碳密度、蒸散、产水和威胁暴露",
                   "需要以矿区斑块而非行政平均值组织监测、预测和修复优先序",
                   "现有成果主要基于遥感分类和模型代理，缺少同尺度长期地面样方与水文观测")
    add_subsection(document, "1.2 研究目标与技术问题", "五期历史重建、2026年多情景预测和生态服务综合评价",
                   "形成了五期LULC、四情景PLUS结果、InVEST三项服务栅格、综合指数、转移矩阵、桑基图和沉陷专题图",
                   "历史变化用于识别趋势，情景差异用于识别政策响应，二者通过统一主网格和分类编码连接",
                   "构建从用户本地数据到地图、统计表、模型记录和报告的一体化成果链",
                   "以统一数据口径、参数记录和空间对照保证结果能够被复核和持续更新")
    add_subsection(document, "1.3 报告结构与证据原则", "论文式章节结构和结果可追溯性",
                   "正文按土地利用、碳储量、水源供给、生境质量、综合服务和情景治理依次展开，每章配置三线表和对应图件",
                   "每个结论均追溯到CSV、GeoTIFF或模型清单，数值采用统一单位，图表与文字共用同一统计源",
                   "将‘输入统一’、‘模型运行’和‘成果核查’作为连续的质量控制环节",
                   "参考文献结构用于组织方法，但不复制两篇论文的示例结果或矿区结论")
    three_line_table(document, "表1-1 研究问题、分析对象与主要成果", ["研究问题", "分析对象", "主要成果"], [
        ["土地利用如何变化", "2005—2025年六类LULC", "面积表、转移矩阵、桑基图"],
        ["生态服务如何响应", "碳储量、供水、生境质量", "五期栅格、统计表、时间图"],
        ["2026年情景差异", "ND、UD、EP、RE", "PLUS与InVEST情景结果"],
        ["综合治理如何分区", "AHP-min-max综合指数", "时空格局与治理建议"],
    ])

    # Chapter 2
    heading(document, "2 研究区、数据与技术路线", 1)
    add_subsection(document, "2.1 研究区与矿区空间特征", "皖北六市矿区的区域背景和空间离散特征",
                   "矿区斑块分布于六市范围内，南北向跨度较大，单体矿区之间被城镇、耕地和水系分隔",
                   "离散斑块使同一行政区内的扰动类型、汇水条件和修复阶段存在显著差异",
                   "地图采用完整六市边界作为参照底图，同时以矿区边界限制成果解释范围",
                   "行政边界用于定位而不是作为生态过程边界，水文和生境分析仍受流域及威胁距离控制")
    add_body(document, "研究区位于安徽省北部淮河平原及其南北过渡地带，研究对象覆盖蚌埠、亳州、阜阳、淮北、淮南和宿州六市的煤矿及其沉陷影响范围。矿区边界图层共记录97个矿区要素，其中淮北市49个、宿州市21个、淮南市16个、亳州市6个、阜阳市5个；蚌埠市在当前矿区边界数据中未形成独立矿区面，但仍保留在六市区域底图中。矿区以条带状、组团状斑块分布，部分矿区与村庄、道路、河沟和农田交错，导致沉陷水体、农业复垦和城镇扩张具有明显的空间邻接关系。")
    add_body(document, "研究区的生态过程具有典型的高潜水位采煤区特征：地下开采引起地表沉降，低洼区在降水和地下水补给条件下可能形成长期或季节性积水；积水又会改变原有耕地、草地和建设用地的利用状态。与此同时，煤矸石堆场、工业场地、交通设施和城镇边缘形成持续的人类活动压力带，林草恢复、湿地修复和光伏利用则在部分沉陷区形成新的土地利用组合。因此，研究区不应被理解为均质的行政单元，而应按照“市域定位—矿区斑块—工作面及沉陷影响区—修复工程单元”四个层级进行解释。")
    add_body(document, "为保证跨期比较，所有结果在EPSG:32650投影下统一到30 m主网格，并以五期共同支撑区作为统计边界。共同支撑区面积为201,232.98 hm²；面积统计采用像元计数乘0.09 hm²，碳储量直接汇总InVEST tot_c_cur有效像元值，年产水量分别报告空间平均深度和体积总量。这样的边界设计避免了不同年份有效范围变化造成的总量偏差，也使历史结果与2026年四情景保持同一空间参照。")
    add_subsection(document, "2.2 数据体系与统一预处理", "遥感影像、ROI、边界、驱动因子、沉陷数据和InVEST参数",
                   "输入包括五期影像及ROI、矿区与六市边界、DEM及地形因子、交通水系距离、人口GDP夜光、ERA5-Land气候、土层深度、PAWC、碳密度和沉陷云图",
                   "分类数据以最近邻或多数值方式对齐，连续驱动因子采用双线性重采样，并统一CRS、范围、30 m像元、行列数和NoData",
                   "采用本地优先的数据处理与软件桥接，保留原始数据只读和派生数据可追溯",
                   "部分气候、土壤和威胁参数为公开数据或论文参数的空间代理，需要后续本地校准")
    add_subsection(document, "2.3 模型与综合评价方法", "PLUS、InVEST和AHP-min-max的串联关系",
                   "PLUS以2020—2025年扩张信息和11类驱动因子估计发展潜力，CARS生成2026年四情景；RE额外叠加沉陷深度和工作面证据",
                   "每个LULC结果分别进入Carbon、Annual Water Yield和Habitat Quality，随后在全时期统一范围内归一化并加权",
                   "以中间栅格、参数表、清单和验证摘要保存模型间契约，避免人工转录造成编码和单位错位",
                   "PLUS GUI运行的复现还依赖固定版本、界面配置和随机种子，必须保存运行证据")
    heading(document, "2.3.1 生态服务模型原理与公式", 3)
    add_body(document, "碳储量采用InVEST Carbon的地类碳密度法。像元x所属地类j的四类碳库相加得到像元碳储量，研究区总量对有效像元求和。")
    add_formula(document, "2-1", [
        _m_sub(_m_run("C"), _m_run("x")), _m_run(" = "),
        _m_sub(_m_run("C"), _m_run("above,x")), _m_run(" + "),
        _m_sub(_m_run("C"), _m_run("below,x")), _m_run(" + "),
        _m_sub(_m_run("C"), _m_run("soil,x")), _m_run(" + "),
        _m_sub(_m_run("C"), _m_run("dead,x")), _m_run(" ;   "),
        _m_sub(_m_run("C"), _m_run("total")), _m_run(" = "),
        _m_sub(_m_run("Σ"), _m_run("x")), _m_run(" "), _m_sub(_m_run("C"), _m_run("x")),
    ])
    add_body(document, "InVEST的tot_c_cur已经把碳密度与像元面积折算为像元碳量，统计总量时直接求和；若制作用地类碳密度图，则统一以Mg C/hm²表示并显式记录面积换算因子。")
    add_body(document, "年产水量采用Budyko水量平衡思想。P为年降水，AET为实际蒸散；参考蒸散由Kc和ET0共同确定，土壤有效水分、根系深度和PAWC控制植被可利用水量。")
    add_formula(document, "2-2", [
        _m_run("Y(x) = [1 − "), _m_frac(_m_run("AET(x)"), _m_run("P(x)")),
        _m_run("]P(x)    ;    PET(x) = "), _m_sub(_m_run("K"), _m_run("c")),
        _m_run("(x) · "), _m_sub(_m_run("ET"), _m_run("0")), _m_run("(x)"),
    ])
    add_body(document, "报告同时输出平均产水深（mm）与总量（m³），前者用于空间比较，后者用于区域水量核算。")
    add_body(document, "生境质量按威胁暴露—地类适宜性框架计算。Hj为地类适宜性，Dxj为综合威胁退化度，k为半饱和常数，z为形状参数。道路、铁路、建设活动和矿业扰动通过最大作用距离、权重和敏感性表进入退化度。")
    # Build separate XML nodes for numerator and denominator.  An OMML node
    # cannot be attached to two parents; reusing the same ``d_pow`` object
    # would silently move it out of the numerator when the denominator is
    # appended and produce a malformed-looking fraction in Word.
    d_pow_num = _m_sup(_m_sub(_m_run("D"), _m_run("xj")), _m_run("z"))
    d_pow_den = _m_sup(_m_sub(_m_run("D"), _m_run("xj")), _m_run("z"))
    k_pow = _m_sup(_m_run("k"), _m_run("z"))
    add_formula(document, "2-3", [
        _m_sub(_m_run("Q"), _m_run("xj")), _m_run(" = "), _m_sub(_m_run("H"), _m_run("j")),
        _m_run("[1 − "), _m_frac(d_pow_num, _m_group([d_pow_den, _m_run(" + "), k_pow])), _m_run("]"),
    ])
    add_body(document, "Hj为地类生境适宜性，Dxj为综合威胁退化度；模型输出为0—1指数，空间解释结合地类、威胁和距离衰减共同进行。")
    add_body(document, "综合生态系统服务先进行全时期min-max标准化，再按AHP权重合成。")
    add_formula(document, "2-4", [
        _m_sub(_m_run("z"), _m_run("i")), _m_run(" = "),
        _m_frac(_m_run("xᵢ − min(xᵢ)"), _m_run("max(xᵢ) − min(xᵢ)")), _m_run(" ;   "),
        _m_run("S = 0.636986"), _m_sub(_m_run("z"), _m_run("carbon")),
        _m_run(" + 0.258285"), _m_sub(_m_run("z"), _m_run("water")),
        _m_run(" + 0.104729"), _m_sub(_m_run("z"), _m_run("habitat")),
    ])
    add_formula(document, "2-5", [
        _m_sub(_m_run("λ"), _m_run("max")), _m_run(" = 3.0385;   CI = "),
        _m_frac(_m_run("λmax − 3"), _m_run("3 − 1")), _m_run(" = 0.0193;   CR = "),
        _m_frac(_m_run("CI"), _m_run("RI")), _m_run(" = 0.0332 < 0.10"),
    ])
    add_body(document, "判断矩阵满足一致性要求，综合指数表达统一参数和共同范围下的相对状态，正文同时保留三项单项服务结果。")
    three_line_table(document, "表2-1 主要数据及其用途", ["数据类别", "代表数据", "用途", "处理规则"], [
        ["遥感与ROI", "五期影像、逐期ROI", "监督分类", "分类编码为1—6"],
        ["地形与区位", "DEM、坡度、坡向、交通水系距离", "PLUS驱动", "连续变量双线性"],
        ["社会经济", "人口、GDP、夜间灯光", "人类活动驱动", "统一到30 m"],
        ["气候土壤", "降水、ET0、土层深度、PAWC", "年产水量", "单位和高程基准核对"],
        ["沉陷证据", "沉陷云图、工作面", "RE情景及积水分析", "正值向下，限制插值范围"],
        ["生态参数", "碳密度、Kc、根系深度、敏感性", "InVEST", "用户参数优先"],
    ])
    # Prefer the editable VisioMaster export when it is available; the PIL
    # fallback keeps the report reproducible on machines without Visio.
    workflow_path = first_existing([
        assets / "图_技术路线流程图_VisioMaster_v4.png",
        assets / "图_技术路线流程图_中英混排.png",
    ])
    if workflow_path is None:
        workflow_path = assets / "图_技术路线流程图_中英混排.png"
        build_workflow_figure(workflow_path)
    add_figure(document, workflow_path, "图2-1 皖北六市矿区生态系统服务评估技术路线（参考两篇论文结构重绘）", width_cm=15.7)
    add_body(document, "技术路线按“数据准备—分类与统一—土地利用模拟—生态服务计算—综合评价与输出”五个环节组织。每个环节均保存输入清单、处理中间栅格和输出摘要：分类环节保留逐期ROI、类别编码和结果核查记录；统一环节以共同支撑区控制统计边界；PLUS环节把四情景需求、转换矩阵和RE沉陷证据写入运行记录；InVEST环节固定生物物理表、气候背景和威胁参数；综合评价环节记录min-max范围、AHP权重和一致性检验。该流程既可用于论文复现，也便于后续只替换某一期影像或某一类参数后增量重算。")

    # Chapter 3
    heading(document, "3 土地利用格局、转移及2026年情景模拟", 1)
    add_subsection(document, "3.1 五期土地利用数量变化", "增强六类土地利用体系的面积演变",
                   f"结果序列中，耕地由2005年的{fmt(data['lulc_area'][2005]['耕地'])} hm²变为2025年的{fmt(data['lulc_area'][2025]['耕地'])} hm²；沉陷积水由{fmt(data['lulc_area'][2005]['沉陷积水'])} hm²变为2025年的{fmt(data['lulc_area'][2025]['沉陷积水'])} hm²。自然水体、林地和草地在相邻期之间还出现幅度较大的变化",
                   "分类面积的期别差异同时可能来自真实土地转换、沉陷水陆变化、修复植被变化，以及跨期影像、训练样本、传感器和类别映射差异",
                   "在跨期一致性复核前，面积表用于定位异常量级和重点空间单元，不把单次跳变直接表述为真实扩张、恢复或政策成效",
                   "自然水体、林地和草地的突变区域结合原始影像、样本和分类参数进行重点判读")
    add_subsection(document, "3.2 土地利用转移路径", "2005—2025年转移矩阵与桑基关系",
                   "转移矩阵中耕地保持量较大，同时存在耕地向林地、草地、建设用地和水体的多向编码转换；自然水体、建设用地等类别也出现较大反向转换",
                   "长期矩阵把多个阶段累积到同一张表中，只能描述两期分类图的编码差异；桑基图用于展示主要流向，矩阵用于核验面积，但二者均不能单独证明某种真实土地转换过程",
                   "治理上可把耕地—沉陷积水、建设用地扩张和林草变化列为影像复核与现场核查的优先对象，再决定是否形成不同的管理措施",
                   "自然水体的反向转换幅度较大，说明水陆边界和沉陷积水判别是转移解释的关键，应结合工作面和水系背景识别主导路径")
    add_subsection(document, "3.3 2026年四情景与资源开采约束", "ND、UD、EP、RE的需求、规则和空间结果",
                   "四情景均已生成30 m LULC输出；ND、UD、EP和RE分别代表当前配置中的趋势、城镇、生态和资源开采规则集合，RE另把沉陷深度与工作面作为核心驱动",
                   "情景输出由发展潜力、邻域效应、转换矩阵、土地需求和约束共同决定；情景名称本身不能证明输出已经符合相应的政策目标",
                   "EP和RE可作为方案对照，但应同时核验各类土地需求、允许转换、刚性保护约束和输出转换矩阵，特别是林地、生境质量与沉陷积水是否与情景含义一致",
                   "需求表、转换矩阵和限制区记录了四情景的配置差异，可据此解释各情景土地利用和服务指标的排序")
    add_subsection(document, "3.4 土地利用变化的阶段特征", "五期土地利用变化的阶段性、集中性与主导转化",
                   f"按当前分类成果，耕地由{fmt(data['lulc_area'][2005]['耕地'])} hm²变为{fmt(data['lulc_area'][2025]['耕地'])} hm²，沉陷积水由{fmt(data['lulc_area'][2005]['沉陷积水'])} hm²变为{fmt(data['lulc_area'][2025]['沉陷积水'])} hm²；但其他类别在相邻年份的增减并不连续",
                   "例如，自然水体在2005—2015年由高值降至极低值，林地在2020年出现高值后又回落。这些量级不宜在未复核前归因于采矿扰动、复垦或自然恢复，而应先追溯影像源、样本、分类器和类别映射是否一致",
                   "耕地—沉陷积水相邻区、自然水体突变区及林草大幅变化区可作为风险管控和影像复核的共同优先单元；通过复核后再判定其是否构成恢复成效或扰动加剧证据",
                   None)
    area_rows = []
    for y in YEARS:
        area_rows.append([str(y)] + [fmt(data["lulc_area"][y][c]) for c in CLASS_ORDER])
    three_line_table(document, "表3-1 2005—2025年土地利用面积（hm²）", ["年份"] + CLASS_ORDER, area_rows,
                     "面积由30 m分类栅格像元数乘0.09 hm²得到；五期采用同一六类编码和共同支撑区进行比较。")
    change_rows = []
    for landuse in CLASS_ORDER:
        initial = data["lulc_area"][2005][landuse]
        final = data["lulc_area"][2025][landuse]
        change_rows.append([landuse, fmt(initial), fmt(final), fmt(final - initial), fmt(pct(final, initial), 2)])
    three_line_table(document, "表3-2 2005—2025年主要地类面积变化", ["地类", "2005年（hm²）", "2025年（hm²）", "变化量（hm²）", "变化率（%）"], change_rows,
                     "变化率为2025年相对2005年的面积变化。")
    three_line_table(document, "表3-3 跨期分类可比性重点检查", ["检查对象", "成果中的量级", "报告处理"], [
        ["自然水体", f"2005年{fmt(data['lulc_area'][2005]['自然水体'])} hm²，2015年{fmt(data['lulc_area'][2015]['自然水体'])} hm²", "不直接解释为水体消失；优先核对影像、样本和沉陷水体分离规则"],
        ["林地", f"2015年{fmt(data['lulc_area'][2015]['林地'])} hm²，2020年{fmt(data['lulc_area'][2020]['林地'])} hm²，2025年{fmt(data['lulc_area'][2025]['林地'])} hm²", "不直接解释为连续恢复或退化；复核传感器、训练样本与类别映射"],
        ["草地", f"2010年{fmt(data['lulc_area'][2010]['草地'])} hm²，2025年{fmt(data['lulc_area'][2025]['草地'])} hm²", "检查低覆盖植被、耕地休耕和裸地的混分"],
        ["沉陷积水", f"2005年{fmt(data['lulc_area'][2005]['沉陷积水'])} hm²，2025年{fmt(data['lulc_area'][2025]['沉陷积水'])} hm²", "叠加工作面、沉陷云图和历史影像后再判断水陆转换"],
    ], "该表不否定原始成果，而是明确后续判读所需的最小一致性检查。")
    for y in YEARS:
        add_figure(document, result / "土地利用分类" / "历史时期" / str(y) / f"{y}_lulc.png",
                   f"图3-{YEARS.index(y)+1} {y}年皖北六市矿区六类土地利用分类")
    add_figure(document, result / "桑基" / "土地利用转化桑基_论文版_2005_2025.png", "图3-7 2005—2025年土地利用转化桑基图")
    add_figure(document, result / "土地利用转移矩阵" / "转移矩阵_2005_2025.png", "图3-8 2005—2025年土地利用转移矩阵")
    add_figure(document, result / "沉陷云图" / "2026" / "2026年皖北六市矿区沉陷云图_A-F.png", "图3-9 2026年皖北六市矿区沉陷空间证据")

    # Chapter 4
    heading(document, "4 碳储量动态与沉陷积水复合碳库", 1)
    add_subsection(document, "4.1 InVEST碳储量方法与单位", "四类碳库的地类赋值和总量核算",
                   "碳密度表覆盖沉陷积水、自然水体、建设用地、耕地、林地和草地，分别配置地上、地下、土壤和死亡有机质碳库",
                   "模型按像元地类调用碳密度，tot_c_cur为单像元总碳储量；密度图可换算为Mg C/hm²，总量则直接对像元值求和",
                   "地类面积和碳密度共同决定总碳储量，高碳密度林地的面积变化会产生较强贡献",
                   "碳密度来自用户表和既有资料，尚缺六市土壤及植被实测分层校准")
    add_subsection(document, "4.2 历史碳储量时空变化", "2005—2025年总量、阶段变化与地类贡献",
                   f"在固定碳密度和各期LULC输入下，总碳储量结果由2005年的{h[2005]['carbon_mg_c']/1e6:.3f}×10^6 Mg C变为2015年的{h[2015]['carbon_mg_c']/1e6:.3f}×10^6 Mg C，2025年为{h[2025]['carbon_mg_c']/1e6:.3f}×10^6 Mg C",
                   "该序列同时受地类面积和固定碳密度控制，尤其会继承跨期LULC分类的量级差异；它可用于比较当前输入条件下的碳库结果，不能直接当作实际碳通量或恢复速率",
                   "保护现有稳定植被和耕地斑块、提高复垦植被稳定性并控制高碳地类向建设和扰动地转换，是待分类可比性确认后仍具有普适性的碳库管理方向",
                   "模型为静态碳密度乘面积框架，不模拟植被年龄、土壤碳恢复时滞和采矿排放")
    add_subsection(document, "4.3 沉陷积水复合碳库", "水体库容、水生植被、底泥面积与碳储量的组合估算",
                   "成果中已形成沉陷积水库容、水生植被覆盖、底泥覆盖和复合碳储量时间序列，并以2026沉陷云图约束潜在积水位置",
                   "沉陷积水碳库由水体溶解和颗粒碳、水生植被碳以及底泥有机碳构成，库容与覆盖面积决定各子库的尺度",
                   "修复方案不能只把积水视为土地损失，还需区分安全水体、污染风险水体和可构建湿地的水体",
                   "水位、底泥厚度、有机碳浓度和植被类型缺少逐水体现场调查，现有复合碳库属于参数化估算")
    carbon_summary: dict[int, dict[str, dict[str, str]]] = {2005: {}, 2025: {}}
    for row in data["carbon_by_class"]:
        year = int(row.get("year", row.get("年份", "0")))
        landuse = row.get("landuse", row.get("地类", ""))
        if year in carbon_summary and landuse in CLASS_ORDER:
            carbon_summary[year][landuse] = row
    class_rows = []
    for landuse in CLASS_ORDER:
        initial = numeric_field(carbon_summary[2005][landuse], "carbon_storage_mg_c", "碳储量_MgC")
        final = numeric_field(carbon_summary[2025][landuse], "carbon_storage_mg_c", "碳储量_MgC")
        share = final / h[2025]["carbon_mg_c"] * 100.0
        class_rows.append([landuse, fmt(initial / 1e6, 3), fmt(final / 1e6, 3), fmt(final / 1e6 - initial / 1e6, 3), fmt(share, 2)])
    add_subsection(document, "4.4 地类碳库与典型矿区差异", "碳储量的地类贡献及典型矿区的阶段性响应",
                   f"2025年耕地碳储量为{numeric_field(carbon_summary[2025]['耕地'], 'carbon_storage_mg_c', '碳储量_MgC')/1e6:.3f}×10^6 Mg C，占研究区总碳储量的{numeric_field(carbon_summary[2025]['耕地'], 'carbon_storage_mg_c', '碳储量_MgC')/h[2025]['carbon_mg_c']*100:.2f}%；林地和草地分别贡献{numeric_field(carbon_summary[2025]['林地'], 'carbon_storage_mg_c', '碳储量_MgC')/1e6:.3f}×10^6 Mg C和{numeric_field(carbon_summary[2025]['草地'], 'carbon_storage_mg_c', '碳储量_MgC')/1e6:.3f}×10^6 Mg C，是耕地以外的重要碳库",
                   "典型矿区序列用于比较各矿区在同一碳密度表和同一处理链下的结果差异。变化幅度既可能与耕地基底、植被、沉陷积水和建设扰动位置有关，也可能继承分期分类差异，不能仅以矿区面积或单期高低判断碳库变化",
                   "对高碳储量矿区，应保持耕地和稳定植被斑块的连续性；对变化明显的矿区，应将复垦地植被、沉陷积水岸带、新增建设用地和分类突变区一并作为重点核查对象",
                   None)
    carbon_rows = []
    for y in YEARS:
        v = h[y]["carbon_mg_c"]
        carbon_rows.append([str(y), fmt(v / 1e6, 3), fmt(pct(v, h[2005]["carbon_mg_c"]), 2)])
    three_line_table(document, "表4-1 2005—2025年矿区碳储量", ["年份", "总碳储量（10^6 Mg C）", "较2005年变化（%）"], carbon_rows,
                     "总量为InVEST tot_c_cur有效像元之和，未再次乘像元面积。")
    three_line_table(document, "表4-2 主要地类碳储量变化", ["地类", "2005年（10^6 Mg C）", "2025年（10^6 Mg C）", "变化量（10^6 Mg C）", "2025年占比（%）"], class_rows)
    mine_lookup: dict[str, dict[int, dict[str, str]]] = {}
    for row in data["typical_mines"]:
        mine_lookup.setdefault(row["mine"], {})[int(row["year"])] = row
    mine_rows = []
    for mine, records in sorted(mine_lookup.items(), key=lambda item: float(item[1][2005]["carbon_t_c"]), reverse=True):
        if 2005 not in records or 2025 not in records:
            continue
        start = float(records[2005]["carbon_t_c"])
        end = float(records[2025]["carbon_t_c"])
        mine_rows.append([mine, records[2025]["city"], fmt(start / 1e3, 2), fmt(end / 1e3, 2), fmt(pct(end, start), 2)])
    if mine_rows:
        three_line_table(document, "表4-3 典型矿区碳储量变化", ["矿区", "所在市", "2005年（10^3 Mg C）", "2025年（10^3 Mg C）", "变化率（%）"], mine_rows)
    for y in YEARS:
        add_figure(document, result / "碳储量" / "历史时期" / str(y) / f"{y}_carbon.png",
                   f"图4-{YEARS.index(y)+1} {y}年皖北六市矿区碳储量空间分布（Mg C）")
    add_figure(document, result / "分析统计图与表" / "沉陷水体与碳储量" / "皖北六市沉陷积水库容与复合碳储量时间序列.png", "图4-6 沉陷积水库容、植被、底泥与复合碳储量变化")
    add_figure(document, result / "分析统计图与表" / "典型矿区碳储量" / "皖北六市典型矿区碳储量变化_2005_2025.png", "图4-7 典型矿区碳储量变化")

    # Chapter 5
    heading(document, "5 年产水量（水源供给潜力）评估", 1)
    add_subsection(document, "5.1 年产水量模型与参数", "降水、参考蒸散、土壤和植被参数的水量平衡",
                   "Annual Water Yield输入包括年降水、ET0、土层深度、PAWC、流域单元、LULC及生物物理表，2026情景沿用2025年气候背景以隔离土地利用差异",
                   "模型以Budyko思想估算实际蒸散和产水深，Kc、根系深度、土壤可利用水分和Z参数共同控制蒸散分配",
                   "在相同气候背景下比较情景，可把差异主要解释为土地利用结构效应；历史期则同时包含气候年际波动",
                   "ET0由现有气候数据代理，Z、Kc和根系深度主要参照资料，缺少流量站校准")
    add_subsection(document, "5.2 历史年产水量变化", "五期年产水深和年产水总量的波动特征",
                   f"五期年产水总量最低为{min(v['water_m3'] for v in h.values())/1e6:.3f}×10^6 m³，最高为{max(v['water_m3'] for v in h.values())/1e6:.3f}×10^6 m³，2025年为{h[2025]['water_m3']/1e6:.3f}×10^6 m³",
                   "年产水量受降水输入和ET0控制较强，土地覆盖、土壤水分和下垫面条件进一步调节栅格结果；该指标表示模型估算的产水潜力，不等同于实际供水量、可利用水量或水质状况",
                   "治理应同时考虑年产水量、调蓄能力和水质，不能把高产水简单等同于高生态价值",
                   "模型没有显式模拟矿井排水、地下水开采、沉陷裂隙渗漏和污染负荷")
    add_subsection(document, "5.3 空间分异与管理含义", "矿区斑块、汇水单元和沉陷水体之间的空间关系",
                   "年产水量图显示不同矿区和流域单元的模型估计存在差异，情景图可识别土地覆盖变化对局地产水潜力的影响",
                   "沉陷区微地形改变可能影响汇水与积水，但栅格年产水量并不等同于水面储量；实际可利用水量还取决于连通性、水质和工程调度",
                   "建议将高产水且水质风险较低的单元纳入生态蓄滞空间，将矿业扰动叠加区纳入重点监测",
                   "当前流域边界由DEM和水系派生，平原区微地形与人工排水系统可能造成分水线不确定性")
    add_subsection(document, "5.4 水源供给的年际波动与调蓄意义", "降水背景下的产水波动及沉陷区水资源利用方向",
                   f"研究期平均产水深由2005年的{h[2005]['water_mean_mm']:.2f} mm降至2010年的{h[2010]['water_mean_mm']:.2f} mm，2015年为{h[2015]['water_mean_mm']:.2f} mm，2020年回升至{h[2020]['water_mean_mm']:.2f} mm，2025年为{h[2025]['water_mean_mm']:.2f} mm。2005年与2020年为相对高值阶段，2010年为低值阶段",
                   "水源供给的振幅显著高于碳储量和生境质量，表明气候年景对产水过程具有直接控制作用；土地利用变化则通过植被蒸散、下垫面入渗和汇流路径调节局地差异。因而，历史时期应将气候波动与土地利用效应分别讨论，不能把全部供水变化归因于矿区开发",
                   "在高产水阶段，应加强沉陷积水区的蓄滞、岸带净化和水质风险管理；在低产水阶段，应优先维护林草覆盖和农田土壤蓄水条件，避免矿业扰动叠加排水工程造成局地水资源紧张",
                   None)
    water_rows = [[str(y), fmt(h[y]["water_mean_mm"], 2), fmt(h[y]["water_m3"] / 1e6, 3), fmt(pct(h[y]["water_m3"], h[2005]["water_m3"]), 2)] for y in YEARS]
    three_line_table(document, "表5-1 2005—2025年年产水量（水源供给潜力）", ["年份", "平均产水深（mm）", "总量（10^6 m³）", "较2005年变化（%）"], water_rows)
    for y in YEARS:
        add_figure(document, result / "水源供给" / "历史时期" / str(y) / f"{y}_water_yield.png",
                   f"图5-{YEARS.index(y)+1} {y}年年产水量（水源供给潜力）空间分布（mm）")

    # Chapter 6
    heading(document, "6 生境质量时空变化", 1)
    add_subsection(document, "6.1 生境质量模型与威胁体系", "地类适宜性、威胁权重、最大作用距离和敏感性",
                   "Habitat Quality以道路、铁路、建设活动和矿业扰动等威胁栅格为基础，结合各地类适宜性及对威胁的敏感性计算退化度和质量指数",
                   "威胁影响随距离衰减，地类敏感性决定同一威胁在不同覆盖上的响应，保护可达性参数用于表达制度约束",
                   "模型适合比较土地利用方案对栖息地格局的相对影响，尤其适合识别高质量斑块、退化边缘和潜在廊道",
                   "威胁权重和敏感性主要由资料与规则生成，缺少本地物种调查和威胁实测")
    add_subsection(document, "6.2 五期生境质量变化", "生境质量均值、空间热点和阶段波动",
                   f"生境质量均值分别为{', '.join(f'{y}年{h[y]["habitat_mean"]:.4f}' for y in YEARS)}，其中2020年最高",
                   "该指数由地类适宜性、威胁权重、最大作用距离和敏感性表共同计算，均值的期别差异既反映输入土地利用与威胁栅格差异，也受参数设定影响；不能把指数升降直接写成物种数量或真实生境状态变化",
                   "应优先保护在多期输入和威胁设置下均保持高值的斑块，并在破碎化严重的矿区之间布设生态连接和缓冲带",
                   "全区均值会掩盖局地退化，且不同年份有效像元范围略有差异")
    add_subsection(document, "6.3 沉陷湿地与生境管理", "沉陷积水、岸带植被和矿业威胁的双重效应",
                   "沉陷积水可形成新的水生和湿地生境，但陡岸、水质污染、孤立水面和持续扰动会降低其实际生态功能",
                   "水体面积增加并不自动带来生境改善，需要结合岸带坡度、植被、连通性和污染源距离判断",
                   "具备修复条件的水体可通过缓坡岸线、挺水植被带和与周边林草斑块连接提升质量",
                   "当前模型没有物种层面占域、繁殖成功率和水质响应数据")
    add_subsection(document, "6.4 生境质量结果的期别差异", "模型指数的低值、高值与输入条件",
                   f"2005、2010、2015、2020和2025年的生境质量均值依次为{h[2005]['habitat_mean']:.4f}、{h[2010]['habitat_mean']:.4f}、{h[2015]['habitat_mean']:.4f}、{h[2020]['habitat_mean']:.4f}和{h[2025]['habitat_mean']:.4f}。其中，2010年形成阶段性低值，2020年达到研究期最高值，2025年出现回落",
                   "这一指数序列说明，模型计算结果并不随单一地类面积线性变化。道路、铁路、建设和采矿扰动的距离衰减，以及林草和水体的空间配置，都会影响指数；高值斑块的连续性比单个斑块面积的短期增加更值得关注",
                   "对2020年后指数回落的矿区，应在变化图中定位低值扩展带，优先核对威胁栅格和LULC输入，再据此安排新增扰动边界管控、林草连接和沉陷水体岸带治理",
                   None)
    habitat_rows = [[str(y), f"{h[y]['habitat_mean']:.4f}", fmt(pct(h[y]["habitat_mean"], h[2005]["habitat_mean"]), 2)] for y in YEARS]
    three_line_table(document, "表6-1 2005—2025年生境质量均值", ["年份", "生境质量指数", "较2005年变化（%）"], habitat_rows)
    for y in YEARS:
        add_figure(document, result / "生境质量" / "历史时期" / str(y) / f"{y}_habitat_quality.png",
                   f"图6-{YEARS.index(y)+1} {y}年生境质量空间分布（0—1）")

    # Chapter 7
    heading(document, "7 综合生态系统服务评价", 1)
    add_subsection(document, "7.1 min-max标准化与AHP权重", "量纲统一、权重求解和一致性检验",
                   "碳储量、水源供给和生境质量在所有声明时期共同确定最小值与最大值，标准化到0—1后，按0.636986、0.258285和0.104729加权",
                   "全时期统一归一化保证年份之间可比较，AHP判断矩阵最大特征根为3.0385，CI为0.0193，CR为0.0332，小于0.1",
                   "权重体现本研究对碳储量的优先关注，但不等同于三项服务的客观经济价值",
                   "权重来自既定判断矩阵，需要开展等权、扰动权重和替代矩阵敏感性分析")
    add_subsection(document, "7.2 综合服务历史变化", "五期综合指数的阶段性和空间格局",
                   f"综合指数均值由2005年的{h[2005]['service_mean']:.4f}降至2010年的{h[2010]['service_mean']:.4f}，2015年回升至{h[2015]['service_mean']:.4f}，2020年达到{h[2020]['service_mean']:.4f}，2025年为{h[2025]['service_mean']:.4f}",
                   "综合指数同时受碳库、年产水量和生境质量的相对位置影响；由于碳储量权重较高，碳储量空间格局对结果贡献更强。指数只反映当前min-max范围和AHP权重下的相对排序，不是实际生态服务货币价值或绝对质量等级",
                   "低值连续区应优先识别主导短板，高值区则强调保护稳定性，避免用统一工程措施覆盖不同成因",
                   "线性加权允许服务间完全补偿，可能掩盖某一服务极低但其他服务较高的风险")
    add_subsection(document, "7.3 协同、权衡与空间分区", "多项服务之间的同向变化、反向变化和空间错位",
                   "历史曲线显示碳储量、生境质量和综合指数并非完全同步，水源供给受气候波动影响更明显",
                   "林草恢复通常有利于碳和生境，但可能通过蒸散增加降低产水；建设和采矿扰动可能提高局地径流却降低生境与碳库",
                   "治理分区应以主导问题和服务组合为依据，而不是只按综合指数高低排序",
                   "五个历史时点不足以支持稳健相关推断，当前协同和权衡仅作描述性分析")
    add_subsection(document, "7.4 综合服务的阶段划分", "综合指数的低值、恢复和调整阶段",
                   f"综合生态系统服务指数在2005年为{h[2005]['service_mean']:.4f}，2010年降至{h[2010]['service_mean']:.4f}，2015年回升至{h[2015]['service_mean']:.4f}，2020年达到{h[2020]['service_mean']:.4f}，2025年回调至{h[2025]['service_mean']:.4f}。据此可将研究期概括为“初始基准—低值波谷—恢复提升—阶段峰值—调整回落”五个阶段",
                   "由于碳储量在综合评价中的权重最高，耕地、林地和草地的面积调整会对指数产生较强影响；而水源供给的年际起伏和生境质量的空间错位则决定了相同碳储量背景下不同矿区的服务差异。综合指数适合作为筛选入口，但具体治理仍需回到单项服务和土地利用变化图进行定位",
                   "建议将长期稳定高值区划为保护维持单元，将低值连续且沉陷、建设或裸地扰动叠加的区域划为优先修复单元，并对综合中值但单项服务短板突出的区域实施针对性补偿措施",
                   None)
    service_rows = [[str(y), f"{h[y]['service_mean']:.4f}", fmt(pct(h[y]["service_mean"], h[2005]["service_mean"]), 2)] for y in YEARS]
    three_line_table(document, "表7-1 AHP权重及一致性", ["指标", "权重", "λmax", "CI", "CR", "结论"], [
        ["碳储量", "0.636986", "3.0385", "0.0193", "0.0332", "通过"],
        ["水源供给", "0.258285", "—", "—", "—", "—"],
        ["生境质量", "0.104729", "—", "—", "—", "—"],
    ], "随机一致性指标RI=0.58，一致性阈值为0.10。")
    three_line_table(document, "表7-2 2005—2025年综合生态系统服务指数", ["年份", "综合指数均值", "较2005年变化（%）"], service_rows)
    add_figure(document, assets / "图_2005至2025年生态系统服务时间变化_单位校正.png", "图7-1 2005—2025年生态系统服务时间变化")
    add_figure(document, result / "时空分布" / "综合生态系统服务时空分布_六期.png", "图7-2 综合生态系统服务时空分布")

    # Chapter 8
    heading(document, "8 2026年四情景生态效应与治理分区", 1)
    s = data["scenario_2026"]
    add_subsection(document, "8.1 四情景生态系统服务比较", "土地利用方案对碳、水、生境和综合指数的影响",
                   f"EP情景碳储量为{s['EP']['carbon_mg_c']/1e6:.3f}×10^6 Mg C、供水量为{s['EP']['water_m3']/1e6:.3f}×10^6 m³、综合指数为{s['EP']['service_mean']:.4f}；RE情景对应值为{s['RE']['carbon_mg_c']/1e6:.3f}×10^6 Mg C、{s['RE']['water_m3']/1e6:.3f}×10^6 m³和{s['RE']['service_mean']:.4f}",
                   "四情景使用相同2025年气候背景，因此年产水量差异主要反映土地利用、土壤和参数框架下的模型结果；碳储量差异由地类面积和碳密度决定",
                   "EP、RE、ND和UD是供比较的规则集合。是否实现生态保护或资源开采目标，不能由名称或单一综合指数判定，需要回查土地需求、转换矩阵、限制区和各单项服务",
                   "情景仅代表规则集合，不是2026年真实发生概率")
    add_subsection(document, "8.2 资源开采情景与沉陷风险", "沉陷深度、工作面与土地利用转换的空间耦合",
                   "RE情景将沉陷深度栅格和工作面范围作为核心驱动，增加沉陷积水和扰动转换的空间可能性，并与其他自然和社会经济驱动共同输入PLUS",
                   "沉陷证据限定在工作面及其影响范围内，避免把最近邻插值无限扩展到整个主网格",
                   "应把高沉陷潜势、耕地集中和生态服务低值重叠区作为监测与治理优先区",
                   "概率积分法结果为地表形变预测，积水形成还受潜水位、地形和排水工程影响")
    add_subsection(document, "8.3 分区治理与实施顺序", "保护、修复、管控和持续监测的组合策略",
                   "成果可支持生态保育区、沉陷湿地修复区、耕地保护与复垦区、矿业扰动严控区及城镇协调区的初步划分",
                   "分区以综合指数为入口，再回看单项服务和土地利用转换，防止高权重指标掩盖水质、生境或耕地风险",
                   "实施上宜先控制新增扰动，再修复连通性和岸带，随后开展碳、水、生境协同监测",
                   "最终边界需要现场核查、产权与工程条件审查以及利益相关者协商")
    add_subsection(document, "8.4 情景土地利用配置及服务排序", "四情景下水体、建设、耕地和植被地类的配置差异",
                   f"EP情景的耕地面积为{s['EP']['lulc_area_ha']['耕地']:.2f} hm²、建设用地为{s['EP']['lulc_area_ha']['建设用地']:.2f} hm²、林地为{s['EP']['lulc_area_ha']['林地']:.2f} hm²；RE情景的耕地为{s['RE']['lulc_area_ha']['耕地']:.2f} hm²，沉陷积水为{s['RE']['lulc_area_ha']['沉陷积水']:.2f} hm²，建设用地为{s['RE']['lulc_area_ha']['建设用地']:.2f} hm²。与EP相比，RE情景耕地减少{fmt(s['RE']['lulc_area_ha']['耕地']-s['EP']['lulc_area_ha']['耕地'])} hm²，沉陷积水增加{fmt(s['RE']['lulc_area_ha']['沉陷积水']-s['EP']['lulc_area_ha']['沉陷积水'])} hm²",
                    f"当前权重下的综合指数排序为{service_order_cn}，生境质量排序为{habitat_order_cn}。这说明不能把情景名称直接等同于“生态最优”，应先审查土地需求、转换规则和生态约束是否实现了预期配置",
                    "单项服务与综合指数并不等价。情景管理应同时查看碳储量、年产水量、生境质量、土地利用面积和转换矩阵，而不是以名称或单一指标决定优劣",
                   None)
    field_panel = assets / "图_典型矿区现场核查照片_三类治理模式.jpg"
    build_field_photo_panel(field_panel, root / "实地图片")
    if field_panel.is_file():
        add_figure(document, field_panel, "图8-1 典型矿区现场核查照片：南湖湿地恢复、临涣农业复垦与潘集新能源利用", width_cm=15.7)
        add_body(document, "现场照片为三类治理模式提供了直观参照。南湖照片显示水面、岸带与游憩设施共同构成湿地恢复型空间；临涣照片呈现沉陷治理与农业复垦并行的地表利用场景；潘集照片体现沉陷水面与漂浮式光伏等新能源利用的组合。照片不用于替代遥感精度验证，但可用于解释情景图中水体、耕地和建设边界的治理含义，并为典型矿区后续调查设计样点：水体区重点记录水深、水质、底泥和水生植被，复垦区重点记录覆土、作物长势和排灌条件，新能源利用区重点记录水面占用、岸带连通和设施扰动。")
    scenario_rows = [[code, SCENARIO_CN[code], fmt(s[code]["carbon_mg_c"] / 1e6, 3), fmt(s[code]["water_m3"] / 1e6, 3),
                      f"{s[code]['habitat_mean']:.4f}", f"{s[code]['service_mean']:.4f}"] for code in SCENARIOS]
    three_line_table(document, "表8-1 2026年四情景生态系统服务比较", ["代码", "情景", "碳储量（10^6 Mg C）", "年产水量（10^6 m³）", "生境质量", "综合指数"], scenario_rows,
                     "四情景年产水量沿用2025年气候背景；结果用于相对比较，不以情景名称替代效果验证。")
    scenario_area_rows = [[landuse] + [fmt(s[code]["lulc_area_ha"][landuse]) for code in SCENARIOS] for landuse in CLASS_ORDER]
    three_line_table(document, "表8-2 2026年四情景土地利用面积（hm²）", ["地类"] + [SCENARIO_CN[code] for code in SCENARIOS], scenario_area_rows)
    add_figure(document, assets / "图_2026年四情景生态系统服务比较_单位校正.png", "图8-2 2026年四情景生态系统服务比较")
    figure_no = 3
    for code in SCENARIOS:
        add_figure(document, result / "土地利用分类" / "PLUS2026" / code / f"2026_lulc_{code}.png",
                   f"图8-{figure_no} 2026年{SCENARIO_CN[code]}情景六类土地利用空间分布")
        figure_no += 1
        add_figure(document, result / "碳储量" / "PLUS2026" / code / f"2026_carbon_{code}.png",
                   f"图8-{figure_no} 2026年{SCENARIO_CN[code]}情景碳储量空间分布（Mg C）")
        figure_no += 1
        add_figure(document, result / "水源供给" / "PLUS2026" / code / f"2026_water_yield_{code}.png",
                   f"图8-{figure_no} 2026年{SCENARIO_CN[code]}情景年产水量空间分布（mm）")
        figure_no += 1
        add_figure(document, result / "生境质量" / "PLUS2026" / code / f"2026_habitat_quality_{code}.png",
                   f"图8-{figure_no} 2026年{SCENARIO_CN[code]}情景生境质量空间分布（0—1）")
        figure_no += 1
        add_figure(document, result / "综合生态系统服务" / "PLUS2026" / code / f"2026_ecosystem_service_{code}.png",
                   f"图8-{figure_no} 2026年{SCENARIO_CN[code]}情景综合生态系统服务空间分布（0—1）")
        figure_no += 1

    # Chapter 9
    heading(document, "9 综合结果判读与治理分区", 1)
    add_subsection(document, "9.1 变化—空间—机制—影响的综合判读", "五期土地利用和生态服务的联动变化",
                   "2005—2025年六类土地利用、碳储量、年产水量、生境质量和综合生态服务已在同一30 m共同支撑区内完成统计，图件、转移矩阵和时间序列共同呈现出阶段性波动",
                   "耕地面积决定区域基底，沉陷积水在工作面和低洼汇水区形成成片斑块，自然水体沿河网及大型水面分布；碳储量随地类碳密度与面积变化响应，年产水量受降水—蒸散背景与土壤持水条件共同调节，生境质量则对威胁距离和斑块连续性敏感",
                   "将三项单项服务与六类土地利用叠加，可识别协同高值区、单项短板区和沉陷—建设复合压力区，把区域均值还原为可执行的矿区单元",
                   "高碳稳定斑块以保护为先，沉陷积水与低生境重叠区优先开展岸带和水质治理，耕地—沉陷边界实施排灌与复垦协同，建设—矿业交叠带严格控制新增扰动")
    add_subsection(document, "9.2 典型矿区分区分析", "南湖、临涣、潘集及代表性矿区的治理模式对照",
                   "南湖表现为湿地恢复型，临涣体现农业复垦与沉陷治理并行，潘集呈现水面与新能源利用组合；三类矿区在水体、耕地、林草斑块和建设威胁的空间组合上具有明显差异",
                   "湿地恢复型矿区的服务提升依赖水体—岸带—林草的连续性，农业复垦型矿区受沉陷深度、覆土和排灌条件制约，新能源利用型矿区还需考虑设施覆盖、岸带通行和生境干扰",
                   "对典型矿区分别建立位置图、六类LULC图、沉陷图、三项服务图和年度曲线，可把区域模型结果转化为工程现场可读的案例卡片",
                   "南湖优先保护水质和岸带连通，临涣优先稳定耕地和排灌系统，潘集优先控制设施占用并设置水面生态缓冲")
    add_subsection(document, "9.3 四情景规划含义", "ND、UD、EP和RE的服务权衡与空间约束",
                   f"四情景综合指数排序为{service_order_cn}，生境质量排序为{habitat_order_cn}；EP碳储量为{s['EP']['carbon_mg_c']/1e6:.3f}×10^6 Mg C、年产水量为{s['EP']['water_m3']/1e6:.3f}×10^6 m³，RE碳储量为{s['RE']['carbon_mg_c']/1e6:.3f}×10^6 Mg C、年产水量为{s['RE']['water_m3']/1e6:.3f}×10^6 m³",
                   "情景差异来自土地需求、转换矩阵、扩张潜力和空间约束的组合：EP强调生态用地配置，UD提高建设扩张压力，RE把沉陷深度和工作面证据纳入资源开采条件，ND提供趋势基线",
                   "情景比选应同时读取土地利用面积、沉陷积水变化和三项服务，而不能以单一综合指数取代空间诊断；四情景共同低值区是刚性治理单元，仅RE显著变差的区域是开采时序和工作面约束单元",
                   "建议以EP和ND作为生态底线与趋势参照，以UD识别城镇扩张压力，以RE开展采动风险压力测试，并将情景差异转化为分区管控和年度监测任务")
    three_line_table(document, "表9-1 综合服务空间分区判读", ["分区类型", "识别特征", "重点对象", "治理方向"], [
        ["协同保护区", "碳储量、生境质量和综合指数均处高值，斑块连续", "稳定林草、水体岸带", "限制新增建设和矿业扰动，维护生态廊道"],
        ["沉陷湿地修复区", "沉陷积水与生境/供水潜力叠加，工作面证据明确", "积水区、岸带、排水口", "水质治理、岸带修复、调蓄与安全利用"],
        ["耕地复垦协同区", "耕地连片且沉陷风险或产水波动较高", "临涣等复垦矿区", "覆土、排灌、土壤碳库与生产功能同步恢复"],
        ["矿业扰动管控区", "RE情景下降明显或建设/道路/矿业威胁叠加", "工作面及交通廊道", "控制开采边界、设置缓冲带、滚动监测"],
    ])
    add_subsection(document, "9.4 图件、统计表与模型结果的衔接", "报告成果链的统一口径和阅读顺序",
                   "当前报告已替换五期及四情景专题图、时间序列图、转移矩阵、桑基图、沉陷图、统计表和综合分析图，所有图件采用统一标题、图例、经纬度、比例尺和单位",
                   "先用面积表和时间图识别变化方向，再用地图定位空间集聚，用转移矩阵和桑基图解释类别流向，最后以碳、水、生境和综合服务图判断生态响应",
                   "同一主网格、同一六类编码和同一单位契约使土地利用—模型—图表—报告形成闭环，读者可从任一统计值追溯到对应栅格和情景目录",
                   "后续更新只需替换新增影像、工作面或参数并沿同一模板重绘，确保历年成果与本次报告保持可比")
    add_subsection(document, "9.5 结果真实性与适用范围", "数据事实、模型结果与现实地表过程的分层判断",
                   f"当前共同支撑区面积为{data['common_boundary_area_ha']:.2f} hm²，五期和四情景均使用30 m主网格；2005—2025年碳储量由{h[2005]['carbon_mg_c']/1e6:.3f}×10^6 Mg C变为{h[2025]['carbon_mg_c']/1e6:.3f}×10^6 Mg C，变化幅度仅{pct(h[2025]['carbon_mg_c'], h[2005]['carbon_mg_c']):.2f}%",
                   f"面积序列同时呈现出建设用地由{data['lulc_area'][2005]['建设用地']:.2f} hm²降至{data['lulc_area'][2025]['建设用地']:.2f} hm²（{pct(data['lulc_area'][2025]['建设用地'], data['lulc_area'][2005]['建设用地']):.1f}%）、沉陷积水由{data['lulc_area'][2005]['沉陷积水']:.2f} hm²增至{data['lulc_area'][2025]['沉陷积水']:.2f} hm²（{pct(data['lulc_area'][2025]['沉陷积水'], data['lulc_area'][2005]['沉陷积水']):.1f}%）以及林地、草地从低基数向较高值变化的特征；这更适合解释为六类分类栅格中的表观重分配，不能直接写成城市建设真实收缩或植被恢复的工程量",
                   "碳储量和综合服务属于统一碳密度、统一气候土壤参数及统一归一化范围下的模型量；水源供给是ERA5-Land降水和ET0驱动的产水潜力；生境质量是威胁距离、权重和地类适宜性共同形成的相对指数；2026年四情景沿用2025年气候背景，表达土地利用规则下的潜在差异而非已经发生的事实。因而，空间位置、边界、单位和图表对应关系可作为高可信的制图事实，五期土地利用和派生服务适合做同口径的相对时空比较；绝对水量、绝对碳库、物种状态和2026年发生概率则应结合参数来源与情景假设解读。论文正文已将建设用地反向变化、沉陷积水扩展和低基数林草变化明确为需要结合原始影像与矿区工程过程解释的现象，不包装成确定性的政策成效")
    three_line_table(document, "表9-2 结果真实性分层与合理用法", ["结果对象", "现实一致性", "可以支持的结论", "不宜直接表述"], [
        ["空间边界、投影、分辨率和图表口径", "高", "定位矿区、比较像元和复现图件", "把共同支撑区当作自然流域或行政统计边界"],
        ["五期六类土地利用面积与转移", "中", "描述分类结果的增减、集聚和流向", "把建设用地减少、林草增加直接等同于现实工程量"],
        ["InVEST碳储量", "中", "比较同一碳密度表下的碳库变化和地类贡献", "直接解释为实测碳通量或碳汇速率"],
        ["年产水量（水源供给潜力）", "中—偏低（绝对量）", "比较气候背景和下垫面条件下的产水潜力", "等同于可供取用水量、库容或实测径流"],
        ["生境质量与综合服务指数", "中—偏低（绝对量）", "识别高低值、威胁叠加和治理优先序", "等同于物种数量、群落完整性或生态价值货币量"],
        ["2026 ND/UD/EP/RE情景", "方案级", "比较规则集合、土地需求和服务权衡", "表述为2026年真实发生概率或确定政策效果"],
    ])

    # Chapter 10
    heading(document, "10 结论与建议", 1)
    add_subsection(document, "10.1 主要结论", "历史变化、情景差异和综合生态效应的归纳",
                   f"五期成果中耕地始终占较大比重，自然水体、林地和草地呈现明显的阶段性变化；在现有LULC和固定碳密度输入下，2025年碳储量较2005年高{pct(h[2025]['carbon_mg_c'], h[2005]['carbon_mg_c']):.2f}%，年产水量期别波动明显，综合指数2010年最低、2020年最高",
                   "碳储量、生境质量和综合指数的差异受土地利用输入与参数共同影响，年产水量还受气候背景控制；AHP权重使碳储量对综合指数的影响更大",
                   "当前配置下EP的综合指数最高而生境质量最低，RE的生境质量最高而碳储量和年产水量相对较低；四情景应作为方案对照，不应根据名称直接判定生态优劣",
                   "结论属于现有数据和参数下的相对评估，不能外推为确定性政策效果")
    add_subsection(document, "10.2 管理建议", "以保护优先、风险管控和适应性修复为核心的行动框架",
                   "建议建立沉陷积水动态台账、稳定高碳斑块保护清单、矿业威胁缓冲带和生态服务年度监测栅格",
                   "先避免新增高风险转换，再针对积水岸带、林草连通和耕地复垦开展差异化修复",
                   "把情景比较纳入矿山年度计划与生态修复方案审查，并用现场监测持续校正模型",
                   "实施需结合水质、工程安全、土地权属和地方发展约束")
    add_subsection(document, "10.3 后续数据与模型完善", "从流程成果走向参数本地化和管理应用",
                   "优先补充典型矿区样地碳密度、流量或水量平衡资料、物种或生境样点以及PLUS历史回代结果",
                   "保持30 m主网格和六类编码稳定，在新增数据后只重算受影响阶段，并同步更新清单和报告表格",
                   "将真实案例匿名化后用于回归测试，逐步形成可安装、可诊断、可恢复和可复现的本地智能体产品",
                   "在参数本地化完成前不追求增加更多模型类型，优先提高数据质量、参数透明度和成果更新效率")
    three_line_table(document, "表10-1 后续工作优先序", ["优先级", "任务", "预期成果"], [
        ["P0", "分类结果复核与PLUS历史回代", "变化像元一致性、FoM与情景稳定性"],
        ["P0", "碳、水、生境现场校准", "本地参数及不确定性区间"],
        ["P1", "情景敏感性与权重扰动", "稳健排序和阈值范围"],
        ["P1", "典型矿区连续监测", "沉陷—水体—服务响应序列"],
        ["P2", "匿名回归案例与自动报告", "可复现发布包"],
    ])

    # Long-form interpretation requested by the user: this is intentionally
    # separated from the reproducible result chapters so it can be shortened
    # for a paper or retained for a project report without changing figures.
    heading(document, "11 研究区案例与方法扩展分析", 1)
    add_extended_analysis(document, h, data, s)

    heading(document, "参考资料说明", 1)
    add_body(document, "本报告的章节组织、方法叙述层次和图表安排参考用户提供的两篇学位论文；论文中的具体研究区数值、精度和结论未作为本项目结果引用。模型方法依据PLUS和InVEST的通用原理及本项目参数、脚本和成果清单表述。正式用于学位论文或期刊投稿时，应按学校或期刊格式补齐原始数据来源、软件版本、模型文献、参数来源及实地调查文献。")

    heading(document, "附录A 成果文件与单位说明", 1)
    three_line_table(document, "表A-1 主要成果类型", ["成果", "主要文件", "单位/范围"], [
        ["土地利用", "LULC_年份_30m_masked.tif", "编码1—6，面积hm²"],
        ["碳储量", "tot_c_cur.tif", "Mg C/像元；总量Mg C"],
        ["水源供给", "wyield_wy_年份.tif", "mm；总量m³"],
        ["生境质量", "quality_年份_30m.tif", "0—1"],
        ["综合服务", "ecosystem_service_年份.tif", "0—1"],
        ["情景土地利用", "PLUS_ND/UD/EP/RE.tif", "编码1—6"],
    ])
    add_body(document, "报告中的面积均以hm²表示，1 hm²等于10,000 m²；30 m像元面积为900 m²，即0.09 hm²。碳储量总量以Mg C表示，Mg与t在质量数值上等价，但为保持InVEST输出语义，正文统一使用Mg C。供水总量以m³或10^6 m³表示，空间图采用年产水深mm。生境质量和综合生态服务为无量纲指数，只有在相同参数、相同归一化范围和相同空间掩膜下才可直接比较。")

    heading(document, "附录B 重点生态问题与分区治理要点", 1)
    appendix_topics = [
        ("一、沉陷积水区的分类治理", "沉陷积水是皖北高潜水位煤矿区最具辨识度的土地利用变化类型。其生态含义取决于水体形成时间、水深、水质、岸坡稳定性、与河沟水系的连通程度及周边土地利用，而不能仅以水面面积判断。对于水质较好、岸带条件稳定且具备连通潜力的积水区，可将其作为湿地化修复和生态蓄滞空间；对于受矿业排水、生活污染或陡岸侵蚀影响较强的积水区，则应优先开展风险排查、岸带整治和污染控制。"),
        ("二、耕地保护与沉陷风险协同", "耕地仍是研究区面积最大的地类，也是区域碳库和粮食生产功能的重要承载空间。耕地转化具有明显的阶段性，既可能来自沉陷积水扩展，也可能来自城镇建设、复垦工程和林草恢复。对工作面影响范围内的连片耕地，应将沉陷预测、地表排水、土壤质量和复垦利用方式纳入同一张管理底图；对于尚未发生明显积水但存在下沉风险的地块，应提前部署排灌工程与种植结构调整。"),
        ("三、林草恢复的质量导向", "林地和草地的增加通常意味着覆盖度提升，但恢复质量还取决于群落稳定性、斑块形状、与周边水体或农田的连接关系以及后续管护。矿山修复中应避免只追求短期绿化面积，而应把耐湿植物配置、土壤基质改良、岸带缓冲和生态廊道建设结合起来。对孤立小斑块，可通过连接带或复垦地植被连续化降低边缘效应；对大面积恢复区，则应关注后续管护不足造成的退化反弹。"),
        ("四、矿业扰动与建设用地边界", "矿业工业场地、道路、居民点和配套设施会形成连续的人类活动压力带，并通过土地占用、噪声、夜间灯光和交通可达性改变周边生态过程。建设用地和工矿活动的管理重点不应局限于斑块内部，而应兼顾其外缘缓冲带。对于与沉陷积水、耕地或高质量生境直接相邻的建设用地，应设置清晰的扩张边界和生态缓冲空间，防止扰动沿道路和河沟方向进一步扩散。"),
        ("五、碳储量提升的空间路径", "区域碳储量提升并不等同于简单扩大某一类土地面积，而取决于不同地类碳密度、面积变化和空间稳定性的共同作用。耕地是现阶段重要碳库，林地具有较高碳密度，草地和湿地恢复则兼具景观连接与碳汇潜力。因而，碳管理宜采用“保护存量高碳斑块—提高复垦地植被质量—控制高碳地类无序转出—监测沉陷水体复合碳库”的组合策略，而不是孤立地追求某一类面积最大化。"),
        ("六、水源供给与水环境目标的区分", "Annual Water Yield反映的是水量过程，不直接代表水质、水资源可利用性或洪涝风险。沉陷区积水和高产水斑块可在一定条件下增强区域调蓄，但也可能因污染、蒸发、藻类繁殖或工程排水而降低生态效益。后续水资源管理应把产水量图与河网、水质监测点、排水工程和取用水需求叠加，区分可利用水资源、生态蓄滞空间和环境风险水体，避免将高产水区简单认定为高生态价值区。"),
        ("七、生境质量与景观连通", "生境质量的空间分异说明，影响生态过程的不仅是地类属性，还包括威胁源距离、斑块破碎程度和景观连通性。稳定高值区应重点保持完整性，避免新的道路、建设或采矿设施切割；中低值但处于林草—水体过渡带的区域，可通过缓冲带、岸带植被和小微湿地串联提升连通；长期低值且扰动叠加的区域，则应作为生态修复和扰动退出的优先对象。"),
        ("八、典型矿区的差异化实施", "典型矿区分析表明，碳储量和生态服务响应并不随矿区面积同比例变化。不同矿区的耕地基底、开采阶段、沉陷积水规模、林草恢复进程和建设强度不同，因而应采用“一矿一策”的实施单元。大型矿区可建立年度栅格监测和分区修复清单，中小矿区可围绕关键沉陷水体、耕地边界和矿业设施周边开展重点整治；跨市矿区还应协调数据口径和治理时序。"),
        ("九、情景结果的规划使用", "四种PLUS情景用于展示不同发展导向下的相对空间响应。自然发展情景适合识别既有趋势延续的风险，城镇发展情景用于识别建设压力可能集中的位置，生态保护情景用于寻找保护和修复的收益空间，资源开采情景则用于预警工作面和沉陷驱动下的耕地、水体与生态服务变化。规划使用时宜把四情景叠加为“稳定区、敏感区、冲突区和优先修复区”四类单元，而非只选取单一情景作为唯一预测。"),
        ("十、综合服务评价的应用边界", "综合生态系统服务指数的作用是把碳储量、水源供给和生境质量放到同一比较框架中，便于识别服务协同区和短板区。实际决策时仍应同时查看单项服务图：例如综合指数较高的区域可能存在水质或地质安全问题，综合指数一般的区域也可能是耕地保护或生态廊道建设的关键节点。综合评价应作为空间筛选和方案比选工具，与矿山安全、国土空间规划、工程条件和投资成本共同构成治理决策依据。"),
    ]
    for topic, paragraph in appendix_topics:
        heading(document, topic, 3)
        add_body(document, paragraph)

    heading(document, "附录C 历史变化与情景响应综合讨论", 1)
    integrated_discussions = [
        ("一、历史结果的总体特征", f"从2005—2025年成果序列看，土地利用分类、沉陷水体、植被覆盖和气候输入共同塑造模型输出。总碳储量在2015年为{h[2015]['carbon_mg_c']/1e6:.3f}×10^6 Mg C，综合服务指数在2020年为{h[2020]['service_mean']:.4f}，二者不在同一年达到最大值；这说明碳库、年产水量和生境质量具有不同的响应节律，阶段性峰值应结合土地利用结构和气候背景综合解释。"),
        ("二、耕地变化的生态含义", f"耕地在五期始终保持主体地位，2025年面积为{data['lulc_area'][2025]['耕地']:.2f} hm²。耕地既承担粮食生产功能，也贡献了区域最大份额的碳储量，因此耕地变化不能仅以“增加或减少”判断优劣：沉陷低洼区耕地转为水体时，重点是保障受影响农田、排灌系统和复垦机会；在适宜复耕地段，则需同时考虑土壤条件、地下水位与后续沉陷风险，避免短期复垦后再次退化。"),
        ("三、林草配置与碳库结果", f"2025年分类结果中林地面积为{data['lulc_area'][2025]['林地']:.2f} hm²、草地面积为{data['lulc_area'][2025]['草地']:.2f} hm²。林地具有较高的单位面积碳密度，草地对地表覆盖、岸带缓冲和景观连通具有重要作用；两类斑块在采矿扰动边缘和沉陷水体周边的连续分布，是碳库与生境协同提升的关键空间单元。"),
        ("四、沉陷积水的双重属性", f"当前分类成果中沉陷积水由2005年的{data['lulc_area'][2005]['沉陷积水']:.2f} hm²变为2025年的{data['lulc_area'][2025]['沉陷积水']:.2f} hm²。其生态含义不能只由面积变化判定，应叠加工作面、沉陷深度、潜水位、地形和历史影像确认水陆转换。经确认的积水区一方面可能压缩耕地利用空间并增加地质环境治理压力，另一方面也可能形成湿地生境、水面调蓄和复合碳库条件。报告将其单独设类，是为了避免与自然水体混合后掩盖矿业活动的空间影响。"),
        ("五、水源供给的气候—下垫面耦合", f"五期总水源供给量在{min(v['water_m3'] for v in h.values())/1e6:.3f}—{max(v['water_m3'] for v in h.values())/1e6:.3f}×10^6 m³之间变化，年际波动远大于碳储量的变化幅度。该特征表明降水和蒸散背景决定年度水量的主趋势，沉陷、复垦、建设和土壤持水条件则通过汇流、入渗和蒸散对局地产水进行再分配；因此应把丰枯水年气候背景与下垫面变化分层讨论。"),
        ("六、生境质量的空间差异", f"生境质量指数从2010年的{h[2010]['habitat_mean']:.4f}变为2020年的{h[2020]['habitat_mean']:.4f}，2025年为{h[2025]['habitat_mean']:.4f}。该结果表示给定土地利用、威胁权重和敏感性参数下的相对空间格局，不是物种监测结果。高值区可作为威胁栅格、林草覆盖、岸带条件和连通性复核的候选区；低值区可作为建设、交通与矿业扰动叠加关系的核查对象，不能只用全区均值推断实际生态恢复。"),
        ("七、综合服务的分区解释", f"综合指数在2010年为{h[2010]['service_mean']:.4f}、2020年为{h[2020]['service_mean']:.4f}、2025年为{h[2025]['service_mean']:.4f}，显示区域生态服务经历了低值、恢复和调整过程。指数的高低反映三项服务的组合，而非某一项服务绝对占优。因此，综合服务高值区应作为稳定性维护对象，低值区应判断是碳库不足、产水偏低还是生境退化所致，再匹配不同的保护、修复或管控手段。"),
        ("八、四情景的管理启示", f"2026年四情景综合指数排序为{service_order_cn}，生境质量排序为{habitat_order_cn}。实际规划可把四种情景作为规则对照：先核查土地需求、转换矩阵和限制区是否符合其定义，再比较空间格局识别不宜开发、冲突和优先修复区域；不以情景名称直接代替生态效益判断。"),
        ("九、典型矿区的跟踪监测", "典型矿区碳储量序列表明，同一研究期内不同矿区的变化方向和幅度存在差异。对信湖、刘店、丁集、顾桥、刘庄、钱营孜、张集和杨柳等代表性矿区，可将土地利用、沉陷深度、水体面积、林草覆盖和生态服务指数建立年度档案，形成“矿区边界—工作面—沉陷影响区—修复工程区”的分级监测体系。这样既能服务区域尺度的趋势判断，也能为具体矿区的生态修复绩效评估提供可追溯依据。"),
        ("十、跨市矿区的协同治理", "皖北六市矿区在行政上分属不同城市，但沉陷、地下水位变化、河网联系、交通走廊和煤炭开发活动具有跨界特征。对于处于市域边缘或水系上下游的矿区，应统一土地利用分类、沉陷深度、碳密度和生态服务指标口径，避免因行政边界不同导致修复标准和监测频次不一致。可探索建立跨市矿区生态服务底图、重点积水区台账和年度变化通报机制，使区域治理从单矿、单市的碎片化处理转向流域—矿区—工程单元协同。"),
        ("十一、成果图件的联合阅读", "单幅土地利用图、碳储量图、水源供给图、生境质量图和综合生态服务图分别回答“地表覆盖如何变化”“碳库分布在哪里”“水量过程如何变化”“生境受何种威胁”和“多项服务如何组合”等问题。报告中的图件应按同一空间位置联合阅读：例如，在沉陷积水扩展且耕地减少的区域，需要同时查看碳储量变化、产水深、生境质量和工作面位置，才能判断其更适合采取排水复垦、湿地构建、岸带修复还是风险管控。"),
        ("十二、从区域评价到工程落地", "区域生态系统服务评价的价值在于提出空间优先序，而不是替代具体工程设计。进入项目实施阶段后，应将优先修复区与矿山地质环境治理、国土空间用途管制、农田保护、水利工程、村庄布局和资金安排衔接，进一步细化到具体地块。对沉陷水体治理，可明确水位控制、岸坡整治、植被恢复和水质监测方案；对耕地复垦，可明确覆土厚度、排灌条件、土壤改良和后续利用方式；对林草恢复，可明确树草配置、管护周期和连通目标。"),
        ("十三、动态更新与适应性管理", "土地利用、工作面推进、沉陷范围和气候背景都在持续变化，生态服务评价应保持动态更新。建议以年度或关键开采节点为周期，更新遥感影像、工作面、沉陷云图和修复工程范围，并在同一30 m主网格上重算面积、碳储量、水源供给、生境质量和综合服务指数。通过比较新增结果与历史序列，可以及时发现沉陷积水扩展、耕地快速转出、林草恢复不稳定或生态服务持续走低的区域，为矿区生态修复方案调整提供依据。"),
    ]
    for topic, paragraph in integrated_discussions:
        heading(document, topic, 3)
        add_body(document, paragraph)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    body_text = "".join(p.text for p in document.paragraphs)
    audit = {
        "report": str(output),
        "paragraph_character_count": len(re.sub(r"\s+", "", body_text)),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
        "validation_status": "pending_validation",
        "notes": [
            "classification independent accuracy metrics are not available",
            "large inter-period LULC jumps are retained as outputs but require harmonized classification review before causal interpretation",
            "PLUS FoM and multi-seed stability are not available",
            "scenario rankings are calculated from the latest complete PLUS/InVEST source set; scenario labels are not treated as verified policy effects",
            "carbon and water-yield units were recalculated from formal outputs",
        ],
    }
    audit_path = output.with_suffix(".audit.json")
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    return audit


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, required=True)
    parser.add_argument("--result", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    root = args.root.expanduser().resolve()
    result = (args.result or root / "结果_重绘").expanduser().resolve()
    output = (args.output or result / "生态系统服务综合报告" / "皖北六市矿区生态系统服务综合评估报告.docx").expanduser().resolve()
    print(json.dumps(build(root, result, output), ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
